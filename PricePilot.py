import streamlit as st
st.set_page_config(page_icon="🎯",layout="wide")
from streamlit_option_menu import option_menu
import os
import pandas as pd
from PIL import Image
import re
import json
from datetime import datetime, timedelta, date
from st_aggrid import AgGrid, GridOptionsBuilder, JsCode, ColumnsAutoSizeMode, GridUpdateMode, DataReturnMode
import openai
from SAPprijs import sap_prices
from Synonyms import synonym_dict
from Articles import article_table
import difflib
from rapidfuzz import process, fuzz
from io import BytesIO
from PyPDF2 import PdfReader
import extract_msg
import pdfplumber
from functools import partial
from database_setup import create_connection, setup_database
import sqlite3
from http.cookies import SimpleCookie
from simple_salesforce import Salesforce, SalesforceLogin
import time
from docx import Document
import xlsxwriter
import getpass
import requests
from requests.auth import HTTPBasicAuth
from requests_ntlm import HttpNtlmAuth 
from office365.runtime.auth.client_credential import ClientCredential
from office365.sharepoint.client_context import ClientContext
from msal import ConfidentialClientApplication
import jwt
import numpy as np
import tempfile
from tempfile import NamedTemporaryFile
import pyodbc
from sqlalchemy import create_engine, text
import urllib
import tempfile
import speech_recognition as sr
import base64
from streamlit_webrtc import webrtc_streamer, WebRtcMode
from pathlib import Path
from striprtf.striprtf import rtf_to_text
import textract
import xlrd
from ProductgroepSynoniemen import Productgroepen_dict

# 🔑 Configuratie
CLIENT_ID = st.secrets.get("SP_CLIENTID")
CLIENT_SECRET = st.secrets.get("SP_CLIENTSECRET")
SP_SITE = st.secrets.get("SP_SITE")
TENANT_ID = st.secrets.get("TENANT_ID")
CSV_PATH = st.secrets.get("SP_CSV_SYN")  # Pad naar TestSynoniem.csv in SharePoint
SP_USERNAME = st.secrets.get("SP_USERNAME")
SP_PASSWORD = st.secrets.get("SP_PASSWORD")

# **Verbinding met Azure SQL Server**
def create_connection():
    server = "vdgbullsaiserver.database.windows.net"
    database = "vdgbullsaidb"
    username = SP_USERNAME
    password = SP_PASSWORD

    conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server};DATABASE={database};Authentication=ActiveDirectoryPassword;UID={username};PWD={password}"
    try:
        conn = pyodbc.connect(conn_str)
        return conn
    except Exception as e:
        st.error(f"Database fout: {e}")
        return None


# Importeer prijsscherpte
if "prijsscherpte_matrix" not in st.session_state:
    # Initialiseer de matrix met standaardwaarden
    st.session_state.prijsscherpte_matrix = pd.DataFrame({
        "Offertebedrag": [0, 5000, 10000, 25000, 50000],  # X-as
        "A": [60, 70, 80, 90, 100],  # Y-as kolommen
        "B": [40, 50, 60, 70, 80],
        "C": [30, 40, 50, 65, 75],
        "D": [10, 25, 45, 60, 65],
    })

st.sidebar.write(f"Laatste update: {time.ctime()}")

# Functie om klantgegevens op te halen uit Salesforce zonder caching
def fetch_salesforce_accounts_direct(sf_connection):
    try:
        # Query voor Salesforce-accounts
        accounts_query = sf_connection.query("""
            SELECT Id, Name, ERP_Number__c
            FROM Account
            WHERE ERP_Number__c != NULL AND Is_Active__c = TRUE
            ORDER BY Name ASC
            LIMIT 6000
        """)
        return accounts_query["records"]
    except Exception as e:
        st.error(f"Fout bij ophalen van Salesforce-accounts: {e}")
        return []

# Salesforce Login Configuratie
SF_USERNAME =  os.getenv("SALESFORCE_USERNAME")
SF_PASSWORD = os.getenv("SALESFORCE_PASSWORD") + os.environ.get("SF_SECURITY_TOKEN")
SF_SECURITY_TOKEN =  os.getenv("SF_SECURITY_TOKEN")
SF_DOMAIN = "test"  # Gebruik 'test' voor Sandbox

if "force_rerun" in st.session_state and st.session_state.force_rerun:
    st.session_state.force_rerun = False  # Zet de trigger uit om oneindige loops te voorkomen
    st.rerun()  # UI herladen zonder dat state verloren gaat


# Verbind met Salesforce
try:
    session_id, instance = SalesforceLogin(
        username=SF_USERNAME,
        password=SF_PASSWORD,
        domain=SF_DOMAIN
    ) #keys in manage app secrets
    sf = Salesforce(instance=instance, session_id=session_id)

except Exception as e:
    st.error(f"Fout bij het verbinden met Salesforce: {e}")

# Haal accounts op als de verbinding geslaagd is
if sf:
    accounts = fetch_salesforce_accounts_direct(sf)
else:
    accounts = []

# Verwerk de accounts als er gegevens beschikbaar zijn
if accounts:
    accounts_df = pd.DataFrame(accounts).drop(columns="attributes", errors="ignore")
    accounts_df.rename(columns={"Name": "Klantnaam", "ERP_Number__c": "Klantnummer"}, inplace=True)
    accounts_df["Klantinfo"] = accounts_df["Klantnummer"] + " - " + accounts_df["Klantnaam"]
else:
    accounts_df = pd.DataFrame(columns=["Klantnaam", "Klantnummer", "Klantinfo"])



# OpenAI API-sleutel instellen
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("OpenAI API-sleutel ontbreekt. Stel de OPENAI_API_KEY omgevingsvariabele in de Streamlit Cloud-instellingen in.")
else:
    openai.api_key = api_key  # Initialize OpenAI ChatCompletion client

pplx_api_key = os.getenv("PERPLEXITY_API_KEY") # Perplexity api key t.b.v. Scout


# Zorg ervoor dat de database bij opstarten correct is
setup_database()

# Hard gecodeerde klantgegevens
customer_data = {
    "111111": {"revenue": "50.000 euro", "size": "D"},
    "222222": {"revenue": "140.000 euro", "size": "B"},
    "333333": {"revenue": "600.000 euro", "size": "A"},
    "100007": {"revenue": "141.000 euro", "size": "B"},
}

# Initialiseer offerte DataFrame en klantnummer in sessiestatus
if "offer_df" not in st.session_state:
    st.session_state.offer_df = pd.DataFrame(columns=["Rijnummer", "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal", "RSP", "SAP Prijs", "Handmatige Prijs", "Min_prijs", "M2 p/s", "M2 totaal", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source"])
if "customer_number" not in st.session_state:
    st.session_state.customer_number = ""
if "loaded_offer_df" not in st.session_state:
    st.session_state.loaded_offer_df = pd.DataFrame(columns=["Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal", "RSP", "M2 p/s", "M2 totaal", "Verkoopprijs"])
if "saved_offers" not in st.session_state:
    st.session_state.saved_offers = pd.DataFrame(columns=["Offertenummer", "Klantnummer", "Eindbedrag", "Datum"])
if "selected_rows" not in st.session_state:
    st.session_state.selected_rows = []

def detect_productgroup_from_text(text):
    text = text.lower().strip()
    for groep, synoniemen in Productgroepen_dict.items():
        for syn in synoniemen:
            if syn.lower() == text:
                return groep  # hoofdnaam van de groep
    return text  # als geen match, gebruik wat er stond


# Converteer article_table naar DataFrame
article_table = pd.DataFrame(article_table)

# Streamlit UI-instellingen
# Maak de tabs aan
tab1, tab2, tab3, tab4, tab5 = st.tabs(["🎯 Offerte Genereren", "⚡ Order-entry", "✨ Beoordeel AI", "⚙️ Beheer","🕵️‍♂️ Scout"])

with tab4:
    st.subheader("Beheer")

    # **Wachtwoordbeveiliging**
    wachtwoord = st.text_input("Voer het wachtwoord in om toegang te krijgen:", type="password")
    if wachtwoord == "Comex25":
        st.success("Toegang verleend tot de beheertab.")

               # Knoppen toevoegen aan de GUI
        col1, col2, col3 = st.columns(3)
        with col1:
        
            with st.expander("1 - 🔍 Bekijk en beheer actieve synoniemen", expanded=False):       
                # **Maak verbinding met de database**
                conn = create_connection()
                if conn:
                    cursor = conn.cursor()
                
                    try:
                        # **Controleer of de tabel 'SynoniemenAI' bestaat**
                        cursor.execute("""
                        SELECT TABLE_NAME FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'SynoniemenAI';
                        """)
                        tabel_bestaat = cursor.fetchone()
                
                        if tabel_bestaat:
                            # **Haal de geaccordeerde synoniemen op**
                            cursor.execute("SELECT Artikelnummer, Synoniem FROM SynoniemenAI")
                            synoniemen_data = cursor.fetchall()
                            
                            # **Haal de kolomnamen op**
                            kolomnamen = [desc[0] for desc in cursor.description]
         
                            
                            # **Controleer of er None-waarden zijn**
                            for rij in synoniemen_data:
                                if None in rij:
                                  pass
                            
                            # **Converteer tuples naar lijsten**
                            synoniemen_data_lijst = [list(rij) for rij in synoniemen_data]
    
                            
                            # **Maak DataFrame aan**
                            synoniemen_df = pd.DataFrame(synoniemen_data_lijst, columns=kolomnamen)
                            
                
                            if not synoniemen_df.empty:
                                # **Zoekbalk toevoegen**
                                zoekterm = st.text_input("🔍 Zoek in synoniemen:", "")
                            
                                # **Filter de DataFrame op basis van de zoekterm**
                                if zoekterm:
                                    synoniemen_df = synoniemen_df[synoniemen_df.astype(str).apply(lambda row: row.str.contains(zoekterm, case=False, na=False)).any(axis=1)]
                            
                                # **Configureer AgGrid voor Synoniemen**
                                gb = GridOptionsBuilder.from_dataframe(synoniemen_df)
                                gb.configure_selection(selection_mode="multiple", use_checkbox=True)
                                gb.configure_default_column(editable=False, filterable=True)  # Filter inschakelen op kolommen
                                grid_options = gb.build()
                            
                                # **Toon de AgGrid met filteropties**
                                response = AgGrid(
                                    synoniemen_df,
                                    gridOptions=grid_options,
                                    update_mode=GridUpdateMode.SELECTION_CHANGED,
                                    fit_columns_on_grid_load=True,
                                    theme="material"
                                )

                
                                # **Geselecteerde rijen ophalen**
                                geselecteerde_rijen = response["selected_rows"]
                
                                if st.button("Verwijder geselecteerde synoniemen"):
                                    if len(geselecteerde_rijen) > 0:
                                        try:
                                            for rij in geselecteerde_rijen:
                                                # Controleer of rij een dictionary is of een tuple/lijst
                                                if isinstance(rij, dict):
                                                    synoniem = rij.get("Synoniem")
                                                    artikelnummer = rij.get("Artikelnummer")
                                                elif isinstance(rij, (tuple, list)) and len(rij) == 2:
                                                    artikelnummer, synoniem = rij  # Pak waarden uit tuple/lijst
                                                else:
                                                    st.warning(f"Ongeldig formaat van rij: {rij}")
                                                    continue
                                
                                                if synoniem and artikelnummer:
                                                    cursor.execute("""
                                                    DELETE FROM SynoniemenAI WHERE Artikelnummer = ? AND Synoniem = ?;
                                                    """, (artikelnummer, synoniem))
                                
                                            conn.commit()
                                            st.success("Geselecteerde synoniemen zijn verwijderd uit 'synoniemen'.")
                                        except Exception as e:
                                            st.error(f"Fout bij verwijderen van synoniemen: {e}")
                                    else:
                                        st.warning("Selecteer minimaal één rij om te verwijderen.")
    
                
                    except Exception as e:
                        st.error(f"Fout bij ophalen van synoniemen: {e}")
                
                    finally:
                        conn.close()
    

    elif wachtwoord:
        st.error("❌ Onjuist wachtwoord. Toegang geweigerd.")

    
    def create_connection():
        server = "vdgbullsaiserver.database.windows.net"
        database = "vdgbullsaidb"
        username = SP_USERNAME
        password = SP_PASSWORD
        driver = "ODBC Driver 17 for SQL Server"
        authentication = "ActiveDirectoryPassword"
    
        params = urllib.parse.quote_plus(
            f"DRIVER={{{driver}}};SERVER={server};DATABASE={database};UID={username};PWD={password};Authentication={authentication}"
        )
    
        conn_str2 = f"mssql+pyodbc:///?odbc_connect={params}"
    
        try:
            engine = create_engine(conn_str2, fast_executemany=True, paramstyle="qmark")
            return engine
        except Exception as e:
            st.error(f"Kan geen verbinding maken met de database: {e}")
            return None
    

    def verwerk_excel(geuploade_bestand):
        if geuploade_bestand is not None:
            try:
                start_time = time.time()
                df = pd.read_excel(geuploade_bestand)
    
                # Controleer of de vereiste kolommen bestaan
                vereiste_kolommen = ["customer number", "product number", "SAP price", "alias customer product"]
                if not all(kolom in df.columns for kolom in vereiste_kolommen):
                    st.error("Excel-bestand mist verplichte kolommen! Zorg dat de kolommen correct zijn.")
                    return
    
                # Hernoem de kolommen naar de database-kolomnamen
                df.rename(columns={
                    "customer number": "customer_number",
                    "product number": "product_number",
                    "SAP price": "SAP_price",
                    "alias customer product": "alias_customer_product"
                }, inplace=True)
    
                engine = create_connection()
                if engine is None:
                    return
                
                # Haal bestaande data op om updates efficiënter te verwerken
                existing_data = pd.read_sql("SELECT alias_customer_product, SAP_price FROM SAP_prijzen", engine)
                existing_data.set_index("alias_customer_product", inplace=True)
                
                df["huidige_prijs"] = df["alias_customer_product"].map(existing_data["SAP_price"].to_dict())
                df["nieuw"] = df["huidige_prijs"].isna()
                df["update_nodig"] = ~df["nieuw"] & (df["SAP_price"] != df["huidige_prijs"])
                
               # **Batch insert voor nieuwe data in kleinere groepen**
                nieuwe_data = df[df["nieuw"]].drop(columns=["nieuw", "update_nodig", "huidige_prijs"])
                batch_size = 11000  # Max 5.000 rijen per batch
                if not nieuwe_data.empty:
                    for i in range(0, len(nieuwe_data), batch_size):
                        nieuwe_data.to_sql("SAP_prijzen", engine, if_exists="append", index=False, method="multi", chunksize=1000)
    
                # **Batch update voor bestaande data**
                update_data = df[df["update_nodig"]]
                if not update_data.empty:
                    update_tuples = list(update_data[["SAP_price", "alias_customer_product"]].itertuples(index=False, name=None))
                    with engine.connect() as connection:
                        statement = text("UPDATE SAP_prijzen SET SAP_price = :SAP_price WHERE alias_customer_product = :alias_customer_product")
               
                        # Directe databaseconnectie via pyodbc
                        conn = pyodbc.connect("DRIVER={ODBC Driver 17 for SQL Server};"
                                              "SERVER=vdgbullsaiserver.database.windows.net;"
                                              "DATABASE=vdgbullsaidb;"
                                              "UID=SP_USERNAME;"
                                              "PWD=SP_PASSWORD;"
                                              "Authentication=ActiveDirectoryPassword")
                        
                        cursor = conn.cursor()
                        
                        # SQL INSERT statement
                        insert_query = """
                        INSERT INTO SAP_prijzen (customer_number, product_number, SAP_price, alias_customer_product) 
                        VALUES (?, ?, ?, ?)
                        """
                        
                        # Zet de data om in tuples voor executemany()
                        data_tuples = [tuple(row) for row in nieuwe_data.itertuples(index=False, name=None)]
                        
                        # **Verwerk in batches om SQL Server limiet te vermijden**
                        batch_size = 11000
                        for i in range(0, len(data_tuples), batch_size):
                            cursor.executemany(insert_query, data_tuples[i:i+batch_size])
                            conn.commit()
                        
                        # Sluit de connectie
                        cursor.close()
                        conn.close()

                        connection.commit()
                end_time = time.time()
                duration = end_time - start_time
                
                st.success(f"✅ Verwerking voltooid in {duration:.2f} seconden! {len(update_data)} prijzen gewijzigd, {len(nieuwe_data)} nieuwe prijzen toegevoegd.")
    
            except Exception as e:
                st.error(f"Fout bij verwerken van Excel-bestand: {e}")
    
    with st.expander("2 - 💲 Upload SAP Prijzen", expanded=False):
        geuploade_bestand = st.file_uploader("Upload prijzen", type=["xlsx"])
        if st.button("📥 Verwerk en sla op in database"):
            verwerk_excel(geuploade_bestand)


        

# Tab 1: Offerte Genereren
with tab1:
    st.subheader("Offerte Genereren")
    
    if st.session_state.offer_df is not None and not st.session_state.offer_df.empty:
        st.title("Offerteoverzicht")




if st.session_state.offer_df is None or st.session_state.offer_df.empty:
    st.session_state.offer_df = pd.DataFrame(columns=["Rijnummer", "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal", "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs", "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source"])


# Omzetting naar numerieke waarden en lege waarden vervangen door 0
st.session_state.offer_df["M2 totaal"] = pd.to_numeric(st.session_state.offer_df["M2 totaal"], errors='coerce').fillna(0)
st.session_state.offer_df["RSP"] = pd.to_numeric(st.session_state.offer_df["RSP"], errors='coerce').fillna(0)
st.session_state.offer_df["Verkoopprijs"] = pd.to_numeric(st.session_state.offer_df["Verkoopprijs"], errors='coerce')

# Offerte Genereren tab
with tab1:
    
    # Voeg een dropdown toe voor prijsbepaling met een breedte-instelling
    col1, _ = st.columns([1, 7])  # Maak kolommen om breedte te beperken
    with col1:
        prijsbepaling_optie = st.selectbox("Prijsbepaling", ["PricePilot logica", "SAP prijs", "RSP"], key="prijsbepaling", help="Selecteer een methode voor prijsbepaling.")

# Offerte Genereren tab
with tab1:
    def bereken_prijs_backend(df):
        if df is None or not isinstance(df, pd.DataFrame):
            st.warning("De DataFrame is leeg of ongeldig. Prijs_backend kan niet worden berekend.")
            return pd.DataFrame()  # Retourneer een lege DataFrame als fallback

        try:
            # Controleer of de DataFrame geldig is
            if not isinstance(df, pd.DataFrame):
                raise ValueError("De input is geen geldige DataFrame.")

            # Zorg ervoor dat kolommen numeriek zijn of bestaan
            for col in ["SAP Prijs", "RSP", "Handmatige Prijs", "Prijskwaliteit"]:
                if col not in df.columns:
                    df[col] = 0  # Voeg de kolom toe als deze niet bestaat

            df["SAP Prijs"] = pd.to_numeric(df["SAP Prijs"], errors="coerce").fillna(0)
            df["RSP"] = pd.to_numeric(df["RSP"], errors="coerce").fillna(0)
            df["Handmatige Prijs"] = pd.to_numeric(df["Handmatige Prijs"], errors="coerce").fillna(0)
            df["Prijskwaliteit"] = pd.to_numeric(df["Prijskwaliteit"], errors="coerce").fillna(100)

            # Functie om Prijs_backend te bepalen op basis van logica
            def bepaal_prijs_backend(row):
                # Controleer of Handmatige Prijs is ingevuld
                if row["Handmatige Prijs"] > 0:
                    return row["Handmatige Prijs"]
                
                # Logica voor SAP Prijs
                elif prijsbepaling_optie == "SAP prijs":
                    return row["SAP Prijs"]
                
                # Logica voor RSP
                elif prijsbepaling_optie == "RSP":
                    rsp_met_kwaliteit = row["RSP"] * (row["Prijskwaliteit"] / 100)
                    return (rsp_met_kwaliteit * 20 // 1 + (1 if (rsp_met_kwaliteit * 20 % 1) > 0 else 0)) / 20
                
                # Logica voor PricePilot
                elif prijsbepaling_optie == "PricePilot logica":
                    # Zorg ervoor dat zowel SAP Prijs als RSP niet 0 zijn
                    if row["SAP Prijs"] > 0 and row["RSP"] > 0:
                        return min(row["SAP Prijs"], row["RSP"])
                    elif row["SAP Prijs"] > 0:
                        return row["SAP Prijs"]  # Gebruik SAP Prijs als RSP 0 is
                    elif row["RSP"] > 0:
                        return row["RSP"]  # Gebruik RSP als SAP Prijs 0 is
                    else:
                        return 0  # Als beide 0 zijn, zet Prijs_backend op 0
                
                # Default naar 0 als niets anders van toepassing is
                return 0

            # Pas de prijsbepaling logica toe op de DataFrame
            df["Prijs_backend"] = df.apply(bepaal_prijs_backend, axis=1)

            # Verkoopprijs is gelijk aan Prijs_backend
            df["Verkoopprijs"] = df["Prijs_backend"]

        except Exception as e:
            st.error(f"Fout bij het berekenen van Prijs_backend: {e}")

        return df








# Controleer en zet kolommen om
for col in ["M2 totaal", "RSP", "Verkoopprijs"]:
    if col not in st.session_state.offer_df.columns:
        st.session_state.offer_df[col] = 0
    st.session_state.offer_df[col] = pd.to_numeric(st.session_state.offer_df[col], errors='coerce').fillna(0)

# Berekeningen uitvoeren
totaal_m2 = st.session_state.offer_df["M2 totaal"].sum()
totaal_bedrag = (st.session_state.offer_df["M2 totaal"] * st.session_state.offer_df["Prijs_backend"]).sum()

# Maak drie kolommen
col1, col2, col3 = st.sidebar.columns(3)

# HTML weergeven in de zijbalk
with col2:
    st.image("BullsAI_logo.png", width=int(30 / 100 * 1024))  # Pas grootte aan (30% van origineel)
    st.sidebar.markdown("---")  # Scheidingslijn voor duidelijkheid  


with tab3:    
    cutoff_value = st.slider(
        "Matchwaarde AI",
        min_value=0.05,
        max_value=1.0,
        value=0.8,  # Standaardwaarde
        step=0.05,  # Stappen in float
        help="Stel matchwaarde in. Hogere waarde betekent strengere matching, 0.8 aanbevolen."
    )
    
    # Bijlagen in mail definiëren
    def detect_relevant_columns(df):
        """
        Detecteert de relevante kolommen (Artikelnaam, Hoogte, Breedte, Aantal) in een DataFrame.
        """
        # Standaardiseer kolomnamen in de DataFrame (trim en lower)
        standardized_columns = {col: col.strip().lower() for col in df.columns}
        
        column_mapping = {
            "Artikelnaam": ["artikelnaam", "artikel", "product", "type", "article", "samenstelling"],
            "Hoogte": ["hoogte", "height", "h"],
            "Breedte": ["breedte", "width", "b"],
            "Aantal": ["aantal", "quantity", "qty", "stuks"]
        }
        detected_columns = {}
    

    
        for key, patterns in column_mapping.items():
            for pattern in patterns:
                match = [original_col for original_col, std_col in standardized_columns.items() if std_col == pattern]
                if match:
                    detected_columns[key] = match[0]
                    break
    
        return detected_columns
    
    
    

        
    # Dynamisch zoeken in de zijbalk
    with st.sidebar:
        search_query = ""
    
        # Filter de resultaten op basis van de invoer
        if not accounts_df.empty and search_query:
            filtered_df = accounts_df[accounts_df["Klantnaam"].str.contains(search_query, case=False, na=False)]
        else:
            filtered_df = accounts_df
    
        # Voeg een lege string toe als eerste optie in de lijst
        klantopties = [""] + filtered_df["Klantinfo"].tolist()
    
        # Toon de selectbox met de lege regel als standaardwaarde
        selected_customer = st.selectbox(
            "Selecteer een klant",
            options=klantopties,
            index=0,  # Hiermee wordt de lege regel standaard geselecteerd
            help="Kies een klant uit de lijst.",
        )
    
        # Afleiden van customer_number als de selectie is gemaakt
        if selected_customer:
            customer_number = selected_customer[:6]  # Haal de eerste 6 tekens uit de selectie
        else:
            customer_number = None
    
        st.session_state.customer_number = str(customer_number) if customer_number else ''
    
        # Klantreferentie invoer
        customer_reference = st.text_input(
            "Klantreferentie",
            value=st.session_state.get("customer_reference", ""),
        )



    offer_amount = totaal_bedrag

# Maak twee kolommen in de sidebar (verhoudingen kunnen aangepast worden)
col1, col2 = st.sidebar.columns([1, 1])

# Linker kolom: Totaalwaarden
with col1:
    st.metric("Totaal m2", f"{totaal_m2:.2f}")
    st.metric("Totaal Bedrag", f"€ {totaal_bedrag:.2f}")

# Rechter kolom: Klantinformatie
with col2:
    if customer_number in customer_data:
        # Haal klantinformatie op
        omzet_klant = customer_data[customer_number]['revenue'].replace("euro", "€").strip()
        klantgrootte = customer_data[customer_number]['size']
        
        # Haal de aangepaste matrix op
        prijsscherpte_matrix = st.session_state.prijsscherpte_matrix
        prijsscherpte = ""

        # Bepaal prijsscherpte op basis van klantgrootte en offertebedrag
        if klantgrootte in prijsscherpte_matrix.columns:
            for index, row in prijsscherpte_matrix.iterrows():
                if offer_amount >= row["Offertebedrag"]:
                    prijsscherpte = row[klantgrootte]
                else:
                    break
        
        # Toon klantinformatie als metrics
        st.metric("Omzet klant", omzet_klant)
        st.metric("Klantgrootte", klantgrootte)
        # st.metric("Prijsscherpte", prijsscherpte) == prijsscherpte nu eruit gehaald, maar kan getoond worden.

# Functie om synoniemen te vervangen in invoertekst
def replace_synonyms(input_text, synonyms):
    for term, synonym in synonyms.items():
        input_text = input_text.replace(term, synonym)
    return input_text

def find_article_details(lookup_article_number, current_productgroup="Alfa", source=None, original_article_number=None):
            


    product_dict = synonym_dict.get(current_productgroup, {})


    if original_article_number is None:
        original_article_number = lookup_article_number  

    # 🔎 Stap 1: Exact match in synonym_dict[productgroup].values()
    if lookup_article_number in product_dict.values():

        
        filtered_articles = article_table[article_table['Material'].astype(str) == str(lookup_article_number)]
        st.write(f"🔍 Gevonden {len(filtered_articles)} rijen in article_table voor materiaal: {lookup_article_number}")
        
        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                lookup_article_number,
                source if source else "synoniem",
                original_article_number,
                None
            )


    # 🔎 Stap 2: Exacte match in synonym_dict[productgroup].keys()
    if lookup_article_number in product_dict.keys():
        matched_article_number = product_dict[lookup_article_number]
        filtered_articles = article_table[article_table['Material'].astype(str) == str(matched_article_number)]

        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                matched_article_number,
                source if source else "synoniem",
                original_article_number,
                None
            )


    # 🔎 Stap 3: Fuzzy match met RapidFuzz
    closest_match = process.extractOne(lookup_article_number, product_dict.keys(), scorer=fuzz.ratio, score_cutoff=cutoff_value * 100)
    if closest_match:
        best_match = closest_match[0]
        matched_article_number = product_dict[best_match]


        filtered_articles = article_table[article_table['Material'].astype(str) == str(matched_article_number)]

        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                matched_article_number,
                source if source else "interpretatie",
                original_article_number,
                best_match
            )
  
    # 🔎 Stap 4: Fuzzy match met difflib
    closest_matches = difflib.get_close_matches(lookup_article_number, product_dict.keys(), n=1, cutoff=cutoff_value)
    if closest_matches:
        best_match = closest_matches[0]
        matched_article_number = product_dict[best_match]
        filtered_articles = article_table[article_table['Material'].astype(str) == str(matched_article_number)]


        if not filtered_articles.empty:
            return (
                filtered_articles.iloc[0]['Description'],
                filtered_articles.iloc[0]['Min_prijs'],
                filtered_articles.iloc[0]['Max_prijs'],
                matched_article_number,
                source if source else "interpretatie",
                original_article_number,
                best_match
            )


    # ❌ Stap 5: Geen match
    return (
        lookup_article_number,
        None,
        None,
        '1000000',
        source if source else "niet gevonden",
        original_article_number,
        None
    )




# Werkt de artikelnummer bij in de DataFrame op basis van de ingevulde artikelnaam. Gebruikt fuzzy matching om de beste overeenkomst te vinden.
def update_article_numbers_from_names(df, article_table, cutoff_value = cutoff_value):
    if df.empty or article_table.empty:
        return df  # Return ongeldige invoer

    for index, row in df.iterrows():
        artikelnaam = row.get("Artikelnaam", "").strip()

        # Alleen bijwerken als er een naam is en de artikelnummer ontbreekt of een slechte match is
        if artikelnaam and (pd.isna(row.get("Artikelnummer")) or row["Source"] in ["niet gevonden", "interpretatie", "GPT"]):

            # Zoek de beste match met fuzzy matching
            best_match = process.extractOne(artikelnaam, article_table["Description"], scorer=fuzz.ratio, score_cutoff=cutoff_value * 100)

            if best_match:
                best_article_name, score, match_index = best_match
                matched_article_number = article_table.iloc[match_index]["Material"]

                df.at[index, "Artikelnummer"] = matched_article_number
                df.at[index, "Source"] = "interpretatie"  # Markeer als fuzzy match
                df.at[index, "fuzzy_match"] = best_article_name  # Voeg fuzzy match kolom toe
            else:
                df.at[index, "Source"] = "niet gevonden"  # Geen match gevonden

    return df



# Functie om aanbevolen prijs te berekenen
def calculate_recommended_price(min_price, max_price, prijsscherpte):
    if min_price is not None and max_price is not None and prijsscherpte != "":
        return min_price + ((max_price - min_price) * (100 - prijsscherpte) / 100)
    return None


# Functie om m2 per stuk te berekenen
def calculate_m2_per_piece(width, height):
    if width and height:
        width_m = int(width) / 1000
        height_m = int(height) / 1000
        m2 = max(width_m * height_m, 0.65)
        return m2
    return None

# Functie om determine_spacer waarde te bepalen uit samenstellingstekst
def determine_spacer(term, default_value="15 - alu"):
    if term and isinstance(term, str):

        # Regex: Zoek alleen naar getallen tussen `-` en `-`
        matches = re.findall(r'-(\d+)-', term)
        values = list(map(int, matches))  # Converteer de gevonden matches naar integers

        # Controleer of er minimaal één getal is gevonden
        if len(values) >= 1:
            spacer_value = values[0]  # Pak de eerste gevonden waarde

            # Controleer of de waarde binnen de juiste range ligt
            if 3 < spacer_value < 30:
                if any(keyword in term.lower() for keyword in ["we", "warmedge", "warm edge"]):
                    result = f"{spacer_value} - warm edge"
                else:
                    result = f"{spacer_value} - alu"
                return result
    return default_value


# Voorbeeld van hoe de waarde wordt opgeslagen in de state
def update_spacer_state(user_input, app_state):
    selected_spacer = determine_spacer(user_input)
    app_state["spacer"] = selected_spacer


# Functie om bestaande spacers niet te overschrijven bij updates
def preserve_existing_spacers(df):
    for index, row in df.iterrows():
        if pd.notna(row.get("Spacer")):
            continue  # Behoud bestaande waarde
        # Alleen waarden aanpassen als deze niet bestaan of leeg zijn
        df.at[index, "Spacer"] = determine_spacer(row.get("Spacer", "15 - alu"))
    return df

# Genereer een mapping van artikelnamen naar artikelnummers
article_mapping = article_table.set_index("Description")["Material"].to_dict()

def update_offer_data(df):
    for index, row in df.iterrows():

        # Stap 1: Oppervlakteberekening
        if pd.notna(row['Breedte']) and pd.notna(row['Hoogte']):
            df.at[index, 'M2 p/s'] = calculate_m2_per_piece(row['Breedte'], row['Hoogte'])

        if pd.notna(row['Aantal']) and pd.notna(df.at[index, 'M2 p/s']):
            df.at[index, 'M2 totaal'] = float(row['Aantal']) * float(str(df.at[index, 'M2 p/s']).split()[0].replace(',', '.'))

        # Stap 2: Artikelinformatie ophalen (alleen als Artikelnummer aanwezig is)
        if pd.notna(row['Artikelnummer']):

            # Voorkom dat we opnieuw op fallback '1000000' zoeken → pak dan originele invoer
            if row['Artikelnummer'] == '1000000' and row.get('original_article_number'):
                lookup_value = row['original_article_number']
            else:
                lookup_value = row['Artikelnummer']

            # Alleen lookup als 'Source' leeg of AI/herkenning is
            if pd.isna(row.get('Source')) or row['Source'] in ['niet gevonden', 'GPT']:
                current_pg = st.session_state.get('current_productgroup', 'Alfa')
                
                description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(
                    lookup_value,
                    current_productgroup=current_pg,
                    original_article_number=row.get('original_article_number') or lookup_value
                )


                # Alleen overschrijven als artikelnaam leeg is of fallback is
                if description and (pd.isna(row.get('Artikelnaam')) or row['Artikelnaam'] == '1000000'):
                    df.at[index, 'Artikelnaam'] = description

                if min_price is not None and max_price is not None:
                    df.at[index, 'Min_prijs'] = min_price
                    df.at[index, 'Max_prijs'] = max_price

                if source:
                    df.at[index, 'Source'] = source
                if original_article_number:
                    df.at[index, 'original_article_number'] = original_article_number
                if fuzzy_match:
                    df.at[index, 'fuzzy_match'] = fuzzy_match

            # Stap 3: SAP Prijs ophalen
            if st.session_state.customer_number in sap_prices:
                sap_prijs = sap_prices[st.session_state.customer_number].get(row['Artikelnummer'], None)
                df.at[index, 'SAP Prijs'] = sap_prijs if sap_prijs else None
            else:
                df.at[index, 'SAP Prijs'] = None

    df = bereken_prijs_backend(df)
    return df



# Functie om de RSP voor alle regels te updaten
def update_rsp_for_all_rows(df, prijsscherpte):
    if prijsscherpte:

        def calculate_rsp(row):
            min_price = row.get('Min_prijs', None)
            max_price = row.get('Max_prijs', None)
            if pd.notna(min_price) and pd.notna(max_price):
                rsp_value = calculate_recommended_price(min_price, max_price, prijsscherpte)
                return round(rsp_value * 20) / 20
            return row.get('RSP', None)

        df['RSP'] = df.apply(calculate_rsp, axis=1)

        # Pas backend-berekeningen toe
        df = bereken_prijs_backend(df)

    return df




# Functie om Prijs_backend te updaten na wijzigingen
def update_prijs_backend():
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

def reset_rijnummers(df):
    if not df.empty:
        df['Rijnummer'] = range(1, len(df) + 1)
    return df


# JavaScript-code voor conditionele opmaak
cell_style_js = JsCode("""
function(params) {
    if (params.colDef.field === "RSP" && params.data.Prijs_backend === params.data.RSP) {
        return {'backgroundColor': '#DFFFD6', 'fontWeight': 'bold'};  // Lichtgroen met vetgedrukte letters
    } else if (params.colDef.field === "SAP Prijs" && params.data.Prijs_backend === params.data["SAP Prijs"]) {
        return {'backgroundColor': '#DFFFD6', 'fontWeight': 'bold'};  // Lichtgroen met vetgedrukte letters
    } else if (params.colDef.field === "Verkoopprijs" && params.data.Prijs_backend === params.data.Verkoopprijs) {
        return {'backgroundColor': '#DFFFD6', 'fontWeight': 'bold'};  // Lichtgroen met vetgedrukte letters
    } else if (params.colDef.field !== "Verkoopprijs") {
        return {'backgroundColor': '#e0e0e0'};  // Grijs voor alle andere cellen
    }
    return null;
}
""")

# Voeg een cell renderer toe om de stericoon weer te geven
cell_renderer_js = JsCode("""
function(params) {
    if (params.data.Source === "interpretatie" || params.data.Source === "GPT") {
        return `✨ ${params.value}`;  // Voeg stericoon toe vóór de waarde
    }
    return params.value;  // Toon de originele waarde
}
""")


def save_changes(df):
    st.session_state.offer_df = df
    st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)
    st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, st.session_state.get('prijsscherpte', ''))

# Offerte Genereren tab
with tab1:
    # Voeg een veld toe voor prijskwaliteit als RSP wordt gekozen met beperkte breedte
    if prijsbepaling_optie == "RSP":
        col1, _ = st.columns([1, 10])
        with col1:
            prijskwaliteit = st.number_input("Prijskwaliteit (%)", min_value=0, max_value=200, value=100, key="prijskwaliteit")
        st.session_state.offer_df["Prijskwaliteit"] = prijskwaliteit

    # Altijd de logica via de functie bereken_prijs_backend toepassen
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

# JavaScript code voor het opslaan van wijzigingen
js_update_code = JsCode('''
function onCellEditingStopped(params) {
    // Opslaan van gewijzigde data na het bewerken van een cel
    let updatedRow = params.node.data;

    // Zorg ervoor dat wijzigingen worden doorgevoerd in de grid
    params.api.applyTransaction({ update: [updatedRow] });
}
''')


# Maak grid-opties aan voor AgGrid met gebruik van een "select all" checkbox in de header
gb = GridOptionsBuilder.from_dataframe(st.session_state.offer_df)
gb.configure_default_column(flex=1, minWidth=50, editable=True)
gb.configure_column("Spacer", editable=True, cellEditor='agSelectCellEditor', cellEditorParams={"values": ["4 - alu", "6 - alu", "7 - alu", "8 - alu", "9 - alu", "10 - alu", "12 - alu", "13 - alu", "14 - alu", "15 - alu", "16 - alu", "18 - alu", "20 - alu", "24 - alu", "10 - warm edge", "12 - warm edge", "14 - warm edge", "15 - warm edge", "16 - warm edge", "18 - warm edge", "20 - warm edge", "24 - warm edge"]})
gb.configure_column("Rijnummer", type=["numericColumn"], editable=False, cellStyle={"backgroundColor": "#e0e0e0"}, cellRenderer=cell_renderer_js)
gb.configure_column(
    "Artikelnaam",
    editable=True,
    cellEditor="agRichSelectCellEditor",
    cellEditorParams={
        "values": list(article_mapping.keys()),  # De mogelijke waarden
        "searchable": True,  # Laat je typen in de dropdown
        "suppressKeyboardEvent": False  # Zorgt dat je kunt typen zonder dat de dropdown sluit
    },
    width=600
)
gb.configure_column("Offertenummer", hide=True)
gb.configure_column("Prijs_backend", hide=True)
gb.configure_column("Min_prijs", hide=True)
gb.configure_column("Artikelnummer", hide=False)
gb.configure_column("Prijskwaliteit", hide=True)
gb.configure_column("Max_prijs", hide=True)
gb.configure_column("Handmatige Prijs", editable=True, type=["numericColumn"])
gb.configure_column("Breedte", editable=True, type=["numericColumn"])
gb.configure_column("Hoogte", editable=True, type=["numericColumn"])
gb.configure_column("Aantal", editable=True, type=["numericColumn"])
gb.configure_column("RSP", editable=False, type=["numericColumn"], valueFormatter="x.toFixed(2)", cellStyle=cell_style_js)
gb.configure_column("Verkoopprijs", editable=True, type=["numericColumn"], cellStyle=cell_style_js, valueFormatter="x.toFixed(2)")
gb.configure_column("M2 p/s", editable=False, type=["numericColumn"], cellStyle={"backgroundColor": "#e0e0e0"}, valueFormatter="x.toFixed(2)")
gb.configure_column("M2 totaal", editable=False, type=["numericColumn"], cellStyle={"backgroundColor": "#e0e0e0"}, valueFormatter="x.toFixed(2)")
gb.configure_column("SAP Prijs", editable=False, type=["numericColumn"], valueFormatter="x.toFixed(2)", cellStyle=cell_style_js)
gb.configure_column("Source", hide=True)
gb.configure_column("fuzzy_match", hide=True)
gb.configure_column("original_article_number", hide=True)


# Configuratie voor selectie, inclusief checkbox in de header voor "select all"
gb.configure_selection(
    selection_mode='multiple',
    use_checkbox=True,
    header_checkbox=True  # Voeg een selectievakje in de header toe
)

# Voeg de JavaScript code toe aan de grid-opties
gb.configure_grid_options(onCellEditingStopped=js_update_code)

# Overige configuratie van de grid
gb.configure_grid_options(domLayout='normal', rowHeight=23)  # Dit zorgt ervoor dat scrollen mogelijk is

# Voeg een JavaScript event listener toe voor updates bij het indrukken van Enter
js_update_code = JsCode('''
function onCellValueChanged(params) {
    let rowNode = params.node;
    let data = rowNode.data;

    // Zorg ervoor dat wijzigingen direct worden toegepast
    params.api.applyTransaction({ update: [data] });

    // Forceer visuele update
    params.api.refreshCells({ force: true });

    // Luister naar de Enter-toets
    document.addEventListener('keydown', function(event) {
        if (event.key === 'Enter') {
            // Ververs de grid wanneer Enter wordt ingedrukt
            params.api.redrawRows();
        }
    });
}
''')
gb.configure_grid_options(onCellValueChanged=js_update_code)

# Bouw grid-opties
grid_options = gb.build()

# Offerte Genereren tab
with tab1:

    # Toon de AG Grid met het material-thema
    edited_df_response = AgGrid(
        st.session_state.offer_df,
        gridOptions=grid_options,
        theme='alpine',
        fit_columns_on_grid_load=True,
        enable_enterprise_modules=True,
        update_mode=GridUpdateMode.SELECTION_CHANGED,
        columns_auto_size_mode=ColumnsAutoSizeMode.FIT_CONTENTS,
        allow_unsafe_jscode=True
    )

    # Update de DataFrame na elke wijziging
    if "data" in edited_df_response:
        updated_df = pd.DataFrame(edited_df_response['data'])
        # Werk de sessiestatus bij met de nieuwe data
        st.session_state.offer_df = updated_df
        # Voer alle benodigde berekeningen uit
        st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
        st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

    
   

    # Verbeterde update_tabel functie
def update_tabel():
    updated_df = pd.DataFrame(edited_df_response['data'])
    st.session_state.offer_df = updated_df
    st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
    st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)

    new_df = st.session_state.offer_df
      
    st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
    st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
    st.session_state["trigger_update"] = True
    st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)

# Offerte Genereren tab
with tab1:
    
    # Knop om de tabel bij te werken
    if st.button("Update tabel"):
        update_tabel()
 
    # Update de DataFrame na elke wijziging
    updated_df = edited_df_response['data']
    save_changes(pd.DataFrame(updated_df))
    
    # Sla de geselecteerde rijen op in sessie status
    selected_rows = edited_df_response.get('selected_rows_id', edited_df_response.get('selected_rows', edited_df_response.get('selected_data', [])))


    # Zorg dat selected_rows geen None of DataFrame is, maar altijd een lijst
    if selected_rows is None or not isinstance(selected_rows, list):
        selected_rows = []
    
    # Als er rijen zijn geselecteerd, zet deze in de sessie state
    if isinstance(selected_rows, list) and len(selected_rows) > 0:
        try:
            st.session_state.selected_rows = [int(r) for r in selected_rows]
        except ValueError:
            st.write("Waarschuwing: Fout bij het converteren van geselecteerde rijen naar indices.")
    else:
        st.session_state.selected_rows = []
    
    def delete_selected_rows(df, selected):
        if selected_rows is not None and len(selected_rows) > 0:
            # Zorg ervoor dat de indices integers zijn
            selected = [int(i) for i in selected]
            st.write("Geselecteerde indices na conversie:", selected)  # Debugging statement
    
            # Verwijder de geselecteerde rijen en reset de index
            new_df = df.drop(index=selected_rows, errors='ignore').reset_index(drop=True)
            return new_df
           
        else:
            return df

   
    # Knoppen toevoegen aan de GUI
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    with col1:
        if st.button("Voeg rij toe"):
            # Voeg een lege rij toe aan het DataFrame
            new_row = pd.DataFrame({
                "Offertenummer": [None], "Artikelnaam": [""], "Artikelnummer": [""], "Spacer": ["15 - alu"], "Breedte": [0], "Hoogte": [0],
                "Aantal": [0], "RSP": [None], "M2 p/s": [0], "M2 totaal": [0], "Min_prijs": [None], "Max_prijs": [None], "Handmatige Prijs": [1000]
            })
            st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_row], ignore_index=True)
            st.session_state.offer_df = bereken_prijs_backend(st.session_state.offer_df)
            # Werk de Rijnummer-kolom bij zodat deze overeenkomt met de index + 1
            st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
            # Vernieuw de AgGrid
            st.rerun()
    
    with col2:
        if st.button("Verwijder rijen (2x klikken)", key='delete_rows_button'):
            # Haal de geselecteerde rijen op in de juiste vorm
            selected = st.session_state.selected_rows
            st.write("Geselecteerde rijen voor verwijdering:", selected)  # Debugging statement
    
            # Verwijder rijen op basis van index
            if len(selected) > 0:
                # Verwijder de rijen uit de DataFrame op basis van de geselecteerde indices
                st.session_state.offer_df = delete_selected_rows(st.session_state.offer_df, selected)
                st.session_state.selected_rows = []  # Reset de geselecteerde rijen na verwijderen
                # Reset de Rijnummer-kolom na verwijderen
                st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
                st.rerun()
            else:
                st.warning("Selecteer eerst rijen om te verwijderen.")
    
        # Zorg dat de update wordt getriggerd na verwijdering
        st.session_state['trigger_update'] = True

  
# Functie om getallen van 1 tot 100 te herkennen
def extract_numbers(text):
    pattern = r'\b(1|[1-9]|[1-9][0-9]|100)\b'
    matches = re.findall(pattern, text)
    return [int(match) for match in matches]

# Functie om woorden naar getallen om te zetten
def word_to_number(word):
    mapping = {
        "een": 1, "twee": 2, "drie": 3, "vier": 4, "vijf": 5, "zes": 6, "zeven": 7, "acht": 8, "negen": 9, "tien": 10, "elf": 11, "twaalf": 12, "dertien": 13, "veertien": 14, "vijftien": 15, "zestien": 16, "zeventien": 17, "achttien": 18, 
        "negentien": 19, "twintig": 20, "eenentwintig": 21, "tweeëntwintig": 22, "drieëntwintig": 23, "vierentwintig": 24, "vijfentwintig": 25, "zesentwintig": 26, "zevenentwintig": 27, "achtentwintig": 28, 
        "negenentwintig": 29, "dertig": 30, "eenendertig": 31, "tweeëndertig": 32, "drieënendertig": 33, "vierendertig": 34, "vijfendertig": 35, "zesendertig": 36, "zevenendertig": 37, "achtendertig": 38, 
        "negenendertig": 39, "veertig": 40, "eenenveertig": 41, "tweeënveertig": 42, "drieënveertig": 43, "vierenveertig": 44, "vijfenveertig": 45, "zesenveertig": 46, "zevenenveertig": 47, "achtenveertig": 48, 
        "negenenveertig": 49, "vijftig": 50, "eenenvijftig": 51, "tweeënvijftig": 52, "drieënvijftig": 53, "vierenvijftig": 54, "vijfenvijftig": 55, "zesenvijftig": 56, "zevenenvijftig": 57, "achtenvijftig": 58, 
        "negenenvijftig": 59, "zestig": 60, "eenenzestig": 61, "tweeënzestig": 62, "drieënzestig": 63, "vierenzestig": 64, "vijfenzestig": 65, "zesenzestig": 66, "zevenenzestig": 67, "achtenzestig": 68, 
        "negenenzestig": 69, "zeventig": 70, "eenenzeventig": 71, "tweeënzeventig": 72, "drieënzeventig": 73, "vierenzeventig": 74, "vijfenzeventig": 75, "zesenzeventig": 76, "zevenenzeventig": 77, "achtenzeventig": 78, 
        "negenenzeventig": 79, "tachtig": 80, "eenentachtig": 81, "tweeëntachtig": 82, "drieëntachtig": 83, "vierentachtig": 84, "vijfentachtig": 85, "zesentachtig": 86, "zevenentachtig": 87, "achtentachtig": 88, 
        "negenentachtig": 89, "negentig": 90, "eenennegentig": 91, "tweeënnegentig": 92, "drieënnegentig": 93, "vierennegentig": 94, "vijfennegentig": 95, "zesennegentig": 96, "zevenennegentig": 97, "achtennegentig": 98, 
        "negenennegentig": 99, "honderd": 100
    }
    return mapping.get(word, None)

# Callback functie voor het verwijderen van geselecteerde rijen
@st.cache_data
def update_dash_table(n_dlt, n_add, data):
    if ctx.triggered_id == "add-row-btn":
        new_row = pd.DataFrame({
            "Offertenummer": [None],
            "Artikelnaam": [""],
            "Artikelnummer": [""],
            "Spacer": [st.session_state.get("last_selected_spacer", "15 - alu")],  # Gebruik de laatst geselecteerde waarde
            "Breedte": [0],
            "Hoogte": [0],
            "Aantal": [0],
            "RSP": [0],
            "M2 p/s": [0],
            "M2 totaal": [0],
            "Min_prijs": [0],
            "Max_prijs": [0],
            "Verkoopprijs": [0]
        })
        df_new_row = pd.DataFrame(new_row)
        updated_table = pd.concat([pd.DataFrame(data), df_new_row])
        return False, updated_table.to_dict("records")

    elif ctx.triggered_id == "delete-row-btn":
        return True, no_update


  
# Functie om het aantal uit tekst te extraheren
def extract_quantity(text):
    # Zoek naar een getal of woord dat voor 'stuks', 'aantal', 'ruiten', 'st', 'keer', of 'x' staat
    unit_matches = re.findall(r'(\d+|twee|drie|vier|vijf|zes|zeven|acht|negen|tien|elf|twaalf|dertien|veertien|vijftien|zestien|zeventien|achttien|negentien|twintig|eenentwintig|tweeëntwintig|drieëntwintig|vierentwintig|vijfentwintig|zesentwintig|zevenentwintig|achtentwintig|negenentwintig|dertig|eenendertig|tweeëndertig|drieëndertig|vierendertig|vijfendertig|zesendertig|zevenendertig|achtendertig|negenendertig|veertig|eenenveertig|tweeënveertig|drieënveertig|vierenveertig|vijfenveertig|zesenveertig|zevenenveertig|achtenveertig|negenenveertig|vijftig|eenenvijftig|tweeënvijftig|drieënvijftig|vierenvijftig|vijfenvijftig|zesenvijftig|zevenenvijftig|achtenvijftig|negenenvijftig|zestig|eenenzestig|tweeënzestig|drieënzestig|vierenzestig|vijfenzestig|zesenzestig|zevenenzestig|achtenzestig|negenenzestig|zeventig|eenenzeventig|tweeënzeventig|drieënzeventig|vierenzeventig|vijfenzeventig|zesenzeventig|zevenenzeventig|achtenzeventig|negenenzeventig|tachtig|eenentachtig|tweeëntachtig|drieëntachtig|vierentachtig|vijfentachtig|zesentachtig|zevenentachtig|achtentachtig|negenentachtig|negentig|eenennegentig|tweeënnegentig|drieënnegentig|vierennegentig|vijfennegentig|zesennegentig|zevenennegentig|achtennegentig|negenennegentig|honderd)\s*(stuks|aantal|ruiten|st|keer|x)\b', text, re.IGNORECASE)

    
    if unit_matches:
        # Als een match gevonden is, zet het om naar een getal
        return word_to_number(unit_matches[0][0]) if unit_matches[0][0].isalpha() else int(unit_matches[0][0])
    
    # Anders zoek naar een getal alleen
    quantity_matches = extract_numbers(text)
    word_matches = re.findall(r'\b(twee|drie|vier|vijf|zes|zeven|acht|negen|tien|elf|twaalf|dertien|veertien|vijftien|zestien|zeventien|achttien|negentien|twintig|eenentwintig|tweeëntwintig|drieëntwintig|vierentwintig|vijfentwintig|zesentwintig|zevenentwintig|achtentwintig|negenentwintig|dertig|eenendertig|tweeëndertig|drieëndertig|vierendertig|vijfendertig|zesendertig|zevenendertig|achtendertig|negenendertig|veertig|eenenveertig|tweeënveertig|drieënveertig|vierenveertig|vijfenveertig|zesenveertig|zevenenveertig|achtenveertig|negenenveertig|vijftig|eenenvijftig|tweeënvijftig|drieënvijftig|vierenvijftig|vijfenvijftig|zesenvijftig|zevenenvijftig|achtenvijftig|negenenvijftig|zestig|eenenzestig|tweeënzestig|drieënzestig|vierenzestig|vijfenzestig|zesenzestig|zevenenzestig|achtenzestig|negenenzestig|zeventig|eenenzeventig|tweeënzeventig|drieënzeventig|vierenzeventig|vijfenzeventig|zesenzeventig|zevenenzeventig|achtenzeventig|negenenzeventig|tachtig|eenentachtig|tweeëntachtig|drieëntachtig|vierentachtig|vijfentachtig|zesentachtig|zevenentachtig|achtentachtig|negenentachtig|negentig|eenennegentig|tweeënnegentig|drieënnegentig|vierennegentig|vijfennegentig|zesennegentig|zevenennegentig|achtennegentig|negenennegentig|honderd)\b', text)

    if word_matches:
        return word_to_number(word_matches[0])  # Neem het eerste gevonden aantal in woorden
    if quantity_matches:
        return quantity_matches[0]  # Neem het eerste gevonden aantal in cijfers
    return None


# Functie om afmetingen (breedte en hoogte) uit tekst te extraheren
def extract_dimensions(text):
    """
    Extraheert breedte en hoogte uit een regel tekst.
    """
    matches = re.findall(r'(\d{3,4})\s*[xX*]\s*(\d{3,4})', text)  # Herken 800x900 of 800*900
    if matches:
        return int(matches[0][0]), int(matches[0][1])  # Eerste geldige combinatie teruggeven
    
    # Alternatieve notatie zoals "700 bij 800"
    matches = re.findall(r'(\d{2,4})\s*bij\s*(\d{2,4})', text, re.IGNORECASE)
    if matches:
        return int(matches[0][0]), int(matches[0][1])

    return None, None


def extract_all_details(line):
    """
    Extraheert het aantal, de afmetingen en het artikelnummer op dynamische wijze uit een regel.
    Zoekt eerst naar een artikelnummer tussen {}, anders gebruikt het de reguliere regex.
    """
    # Stap 1: Extract aantal (kan vóór of na samenstelling staan)
    quantity = extract_quantity(line)

    # Stap 2: Extract breedte en hoogte (kan vóór of na samenstelling staan)
    width, height = extract_dimensions(line)

    # Stap 3: Zoek eerst of er een omschrijving tussen {} staat
    article_number_match = re.search(r'\{([^}]*)\}', line)  # Alles tussen { }
    
    if article_number_match:
        article_number = article_number_match.group(1).strip()  # Verwijder { } en extra spaties
    else:
        # Stap 4: Verwijder het aantal en de afmetingen tijdelijk uit de regel
        clean_line = line
        if quantity:
            clean_line = re.sub(r'\b' + str(quantity) + r'\s*[xX]?\b', '', clean_line)  # Verwijder bijv "4x"
        if width and height:
            clean_line = re.sub(r'\b' + str(width) + r'\s*[xX*]\s*' + str(height) + r'\b', '', clean_line)  # Verwijder bijv "800x800"

        # Stap 5: Gebruik de reguliere regex voor het artikelnummer als er geen {}-omschrijving is
        article_number_match = re.search(r'(\d+(?:[./-]\d+)*(?:[-*#]\d+(?:[./-]\d+)*)+)', line)
        article_number = article_number_match.group(0).strip() if article_number_match else None

    return quantity, width, height, article_number



def handle_gpt_chat():
    if customer_input:
        lines = customer_input.splitlines()
        data = []
        current_article_number = None
        current_productgroup = "Alfa"  # default productgroep

        for line in lines:
            # Detecteer *productgroep* aanduiding zoals *Eclaz One*
            group_match = re.match(r"\*(.+?)\*", line.strip())
            if group_match:
                raw_input_group = group_match.group(1).strip()
                detected_group = detect_productgroup_from_text(raw_input_group)
                current_productgroup = detected_group
                continue

            # Zoek artikelnummer in tekst
            detected_article_number = re.search(r'([A-Za-z0-9/.]+(?:\s*[-/*#]\s*[A-Za-z0-9/.]+)*)', line)
            if detected_article_number:
                current_article_number = detected_article_number.group(0).replace(" ", "")

            # M2-detectie (optioneel)
            m2_match = re.search(r'(\d+)\s*m2.*?(\d+-\d+)|^(\d+-\d+).*?(\d+)\s*m2', line, re.IGNORECASE)

            # Extract standaard details
            quantity, width, height, article_number = extract_all_details(line)

            # Als geen artikelnummer gevonden, gebruik de vorige
            if not article_number and current_article_number:
                article_number = current_article_number
                st.sidebar.info(f"Geen nieuw artikelnummer gevonden, gebruik vorige: {article_number}")

            # Zoek artikelnummer binnen juiste productgroep
            lookup_article_number = synonym_dict.get(current_productgroup, {}).get(article_number, article_number)

            # Verwerking voor m2-regels
            if m2_match:
                if m2_match.group(1):
                    m2_total = int(m2_match.group(1))
                    lookup_article_number = synonym_dict.get(current_productgroup, {}).get(m2_match.group(2), m2_match.group(2))
                else:
                    lookup_article_number = synonym_dict.get(current_productgroup, {}).get(m2_match.group(3), m2_match.group(3))
                    m2_total = int(m2_match.group(4))

                # Zoek artikeldata
                description, min_price, max_price, artikelnummer, source, original_article_number, fuzzy_match = find_article_details(lookup_article_number,current_productgroup=current_productgroup)


                if description:
                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None, description, artikelnummer, None, None, None, None,
                        None, f"{m2_total:.2f}",
                        f"{recommended_price:.2f}" if recommended_price else 0,
                        None, None, min_price, max_price, verkoopprijs,
                        prijs_backend, source, fuzzy_match, original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{lookup_article_number}' niet gevonden in de artikelentabel.")

            # Verwerking als er aantal, breedte & hoogte is
            elif quantity and (width and height):
                description, min_price, max_price, artikelnummer, source, original_article_number, fuzzy_match = find_article_details(lookup_article_number,current_productgroup=current_productgroup)

                if description:
                    spacer = determine_spacer(line)
                    m2_per_piece = round(calculate_m2_per_piece(width, height), 2)
                    m2_total = round(float(quantity) * m2_per_piece, 2)

                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None, description, artikelnummer, spacer, width, height, quantity,
                        f"{m2_per_piece:.2f}", f"{m2_total:.2f}",
                        f"{recommended_price:.2f}" if recommended_price else 0,
                        None, None, min_price, max_price, verkoopprijs,
                        prijs_backend, source, fuzzy_match, original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{lookup_article_number}' niet gevonden in de artikelentabel.")
            else:
                st.sidebar.warning("Regel genegeerd: geen geldige breedte, hoogte of aantal gevonden.")

        # Zet verzamelde data in de offerte-tabel
        if data:
            new_df = pd.DataFrame(data, columns=[
                "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", "Aantal",
                "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs",
                "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend",
                "Source", "fuzzy_match", "original_article_number"
            ])
            new_df.insert(0, 'Rijnummer', new_df.index + 1)

            st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
            st.session_state.offer_df = update_offer_data(st.session_state.offer_df)
            st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
            st.session_state["trigger_update"] = True
            st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
            st.session_state.offer_df = update_article_numbers_from_names(st.session_state.offer_df, article_table)
            st.rerun()
        else:
            st.sidebar.warning("Geen gegevens gevonden om toe te voegen.")

    elif customer_file:
        handle_file_upload(customer_file)
    else:
        st.sidebar.warning("Voer alstublieft tekst in of upload een bestand.")

#  Werkt de SAP Prijs bij op basis van het klantnummer en artikelnummer.
def update_sap_prices(df):   
    for index, row in df.iterrows():
        artikelnummer = row.get('Artikelnummer')
        if artikelnummer and st.session_state.customer_number in sap_prices:
            df.at[index, 'SAP Prijs'] = sap_prices[st.session_state.customer_number].get(artikelnummer, None)
        else:
            df.at[index, 'SAP Prijs'] = None
    df = bereken_prijs_backend(df)  # Herbereken de prijzen
    return df  # Nu staat return binnen de functie
    
# Functie direct uitvoeren en opslaan in sessiestatus
st.session_state.offer_df = update_sap_prices(st.session_state.offer_df)

# Functie voor het verwerken van e-mailinhoud naar offerte
def handle_email_to_offer(email_body):
    if email_body:
        lines = email_body.splitlines()
        data = []
        current_article_number = None
        current_productgroup = "Alfa"  # Pas dit aan als je dynamisch productgroepen wil herkennen

        for line in lines:
            # Optioneel: herken productgroep als die in de tekst staat (zoals *Eclaz One*)
            group_match = re.match(r"\*(.+?)\*", line.strip())
            if group_match:
                current_productgroup = group_match.group(1).strip()
                continue  # Deze regel bevat enkel de groep

            # Detecteer artikelnummer
            detected_article_number = re.search(r'(\d+[./-]?\d*[-*#]\d+[./-]?\d*)', line)
            if detected_article_number:
                current_article_number = detected_article_number.group(0)
                st.sidebar.info(f"Nieuw artikelnummer gevonden: {current_article_number}")

            # Detecteer m2-format
            m2_match = re.search(r'(\d+)\s*m2.*?(\d+-\d+)|^(\d+-\d+).*?(\d+)\s*m2', line, re.IGNORECASE)

            # Extract details zoals aantal, afmetingen en artikelnummer
            quantity, width, height, article_number = extract_all_details(line)

            # Fallback naar eerder gevonden artikelnummer
            if not article_number and current_article_number:
                article_number = current_article_number
                st.sidebar.info(f"Geen nieuw artikelnummer gevonden, gebruik vorige: {article_number}")

            # Gebruik synoniemen mapping met productgroep
            lookup_article_number = synonym_dict.get(current_productgroup, {}).get(article_number, article_number)

            # Verwerking m2-regels
            if m2_match:
                if m2_match.group(1):
                    m2_total = int(m2_match.group(1))
                    article_number = m2_match.group(2)
                else:
                    article_number = m2_match.group(3)
                    m2_total = int(m2_match.group(4))

                lookup_article_number = synonym_dict.get(current_productgroup, {}).get(article_number, article_number)

                description, min_price, max_price, artikelnummer, source, original_article_number, fuzzy_match = find_article_details(
                    lookup_article_number,
                    current_productgroup=current_productgroup
                )

                if description:
                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None, description, artikelnummer, None, None, None, None,
                        None, f"{m2_total:.2f}",
                        f"{recommended_price:.2f}" if recommended_price else 0,
                        None, None, min_price, max_price, verkoopprijs, prijs_backend,
                        source, fuzzy_match, original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{lookup_article_number}' niet gevonden in de artikelentabel.")

            # Verwerking regels met aantal en afmetingen
            elif quantity and (width and height):
                description, min_price, max_price, artikelnummer, source, original_article_number, fuzzy_match = find_article_details(
                    lookup_article_number,
                    current_productgroup=current_productgroup
                )

                if description:
                    spacer = determine_spacer(line)
                    m2_per_piece = round(calculate_m2_per_piece(width, height), 2)
                    m2_total = round(float(quantity) * m2_per_piece, 2)

                    recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
                    verkoopprijs = None
                    prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

                    data.append([
                        None, description, artikelnummer, spacer, width, height, quantity,
                        f"{m2_per_piece:.2f}", f"{m2_total:.2f}",
                        f"{recommended_price:.2f}" if recommended_price else 0,
                        min_price, None, max_price, None, verkoopprijs, prijs_backend,
                        source, fuzzy_match, original_article_number
                    ])
                else:
                    st.sidebar.warning(f"Artikelnummer '{lookup_article_number}' niet gevonden in de artikelentabel.")
            else:
                st.sidebar.warning("Regel genegeerd: geen geldige breedte, hoogte of aantal gevonden.")

        # Zet data in offerteoverzicht
        if data:
            new_df = pd.DataFrame(data, columns=[
                "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", 
                "Aantal", "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs", 
                "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source", "fuzzy_match", "original_article_number"
            ])
            st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
            st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
            st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
            st.rerun()
        else:
            st.sidebar.warning("Geen gegevens gevonden om toe te voegen.")




def handle_mapped_data_to_offer(df):
    """
    Verwerkt de gemapte data en vertaalt deze naar de tabelstructuur voor offertes.
    """
    data = []
    for _, row in df.iterrows():
        description = row["Artikelnaam"]
        height = row["Hoogte"]
        width = row["Breedte"]
        quantity = row["Aantal"]

        # Synoniem lookup en artikelgegevens ophalen
        lookup_article_number = synonym_dict.get(current_productgroup, {}).get(description, description)
        description, min_price, max_price, article_number, source, original_article_number, fuzzy_match = find_article_details(lookup_article_number,current_productgroup=current_productgroup)

        if description:
            recommended_price = calculate_recommended_price(min_price, max_price, prijsscherpte)
            verkoopprijs = None
            prijs_backend = verkoopprijs if verkoopprijs is not None else recommended_price

            m2_per_piece = round(calculate_m2_per_piece(width, height), 2) if width and height else None
            m2_total = round(float(quantity) * m2_per_piece, 2) if m2_per_piece and quantity else None

            data.append([
                None, description, article_number, None, width, height, quantity,
                f"{m2_per_piece:.2f}" if m2_per_piece is not None else None,
                f"{m2_total:.2f}" if m2_total is not None else None,
                f"{recommended_price:.2f}" if recommended_price is not None else 0,
                min_price, None, max_price, None, verkoopprijs, prijs_backend,
                source, fuzzy_match, original_article_number
            ])
        else:
            st.sidebar.warning(f"Artikelnaam '{description}' niet gevonden in de artikelentabel.")

    if data:
        new_df = pd.DataFrame(data, columns=[
            "Offertenummer", "Artikelnaam", "Artikelnummer", "Spacer", "Breedte", "Hoogte", 
            "Aantal", "M2 p/s", "M2 totaal", "RSP", "SAP Prijs", "Handmatige Prijs", 
            "Min_prijs", "Max_prijs", "Verkoopprijs", "Prijs_backend", "Source", "fuzzy_match", "original_article_number"
        ])
        st.session_state.offer_df = pd.concat([st.session_state.offer_df, new_df], ignore_index=True)
        st.session_state.offer_df = update_rsp_for_all_rows(st.session_state.offer_df, prijsscherpte)
        st.session_state.offer_df = reset_rijnummers(st.session_state.offer_df)
        st.rerun()
    else:
        st.sidebar.warning("Geen gegevens gevonden om toe te voegen.")

def remap_and_process(df):
    # Hier kun je logica toevoegen om de achtergehouden regels opnieuw te mappen
    st.write("Her-mapping van achtergehouden regels...")
    return df


def manual_column_mapping(df, detected_columns):
    """
    Biedt de gebruiker een interface om ontbrekende kolommen handmatig te mappen
    en zorgt ervoor dat de kolommen 'Aantal', 'Hoogte' en 'Breedte' numeriek worden gemaakt.
    """
    all_columns = list(df.columns)
    mapped_columns = {k: v for k, v in detected_columns.items() if v in all_columns}

    st.write("Controleer of de kolommen correct zijn gedetecteerd.✨ Indien niet, selecteer de juiste kolom.")

    for key in ["Artikelnaam", "Hoogte", "Breedte", "Aantal"]:
        # Bepaal standaardindex veilig
        try:
            default_index = all_columns.index(mapped_columns[key]) + 1  # +1 vanwege "Geen" als extra optie
        except (KeyError, ValueError):
            default_index = 0

        # Toon selectbox aan gebruiker
        mapped_columns[key] = st.selectbox(
            f"Selecteer kolom voor '{key}'", 
            options=["Geen"] + all_columns,
            index=default_index
        )

    # Filter de mapping om alleen daadwerkelijke selecties te behouden
    mapped_columns = {k: v for k, v in mapped_columns.items() if v != "Geen"}

    # Converteer de kolommen 'Hoogte', 'Breedte' en 'Aantal' naar numeriek als ze zijn geselecteerd
    for key in ["Hoogte", "Breedte", "Aantal"]:
        if key in mapped_columns:
            if df[mapped_columns[key]].dtype not in [np.float64, np.int64]:
                try:
                    df[mapped_columns[key]] = pd.to_numeric(df[mapped_columns[key]], errors="coerce").fillna(0)
                except Exception as e:
                    st.error(f"Fout bij conversie van '{mapped_columns[key]}' naar numeriek: {e}")

    # Toon een waarschuwing voor niet-gemapte kolommen
    for key in ["Artikelnaam", "Hoogte", "Breedte", "Aantal"]:
        if key not in mapped_columns:
            st.warning(f"'{key}' is niet gemapt.")

    return mapped_columns


# Functie voor PDF naar Excel conversie
def pdf_to_excel(pdf_reader, excel_path):
    try:
        with pdfplumber.open(pdf_reader) as pdf:
            writer = pd.ExcelWriter(excel_path, engine='openpyxl')
            has_data = False

            for i, page in enumerate(pdf.pages):
                table = page.extract_table()
                if table and len(table) > 1:
                    headers = table[0] if all(isinstance(h, str) for h in table[0]) else [f"Kolom_{j}" for j in range(len(table[0]))]
                    df = pd.DataFrame(table[1:], columns=headers)
                    if not df.empty:
                        df.to_excel(writer, sheet_name=f"Page_{i+1}", index=False)
                        has_data = True

            if has_data:
                writer.close()
                return excel_path
            else:
                writer.close()
                return None

    except Exception as e:
        pass
        return None




def is_valid_numeric(value, min_value):
    """ Controleert of een waarde numeriek is en groter dan een minimale waarde. """
    try:
        num = float(value)
        return num > min_value
    except (ValueError, TypeError):
        return False

def shift_row_left(row_values, start_index, shift_amount):
    """ Schuift alle waarden rechts van start_index naar links met shift_amount. """
    new_row = row_values.copy()
    new_row[start_index:-shift_amount] = new_row[start_index+shift_amount:]
    new_row[-shift_amount:] = None  # Maak de laatste kolommen leeg na verschuiving
    return new_row

def correct_backlog_rows(df_backlog):
    """
    Corrigeer rijen die in de backlog zitten door de kolommen systematisch naar links te verschuiven
    vanaf de eerste None-waarde.
    """
    corrected_rows = []
    
    for _, row in df_backlog.iterrows():
        row_values = row.values.copy()
        none_index = np.where(pd.isna(row_values))[0]
        
        if len(none_index) > 0:
            none_col = none_index[0]  # Eerste None waarde gevonden
            
            for shift in [1, 2]:  # Probeer 1 en 2 kolommen naar links te schuiven
                corrected_row = shift_row_left(row_values, none_col, shift)
                corrected_series = pd.Series(corrected_row, index=df_backlog.columns)
                
                if (
                    is_valid_numeric(corrected_series["aantal"], 0) and
                    is_valid_numeric(corrected_series["breedte"], 99) and
                    is_valid_numeric(corrected_series["hoogte"], 99)
                ):
                    corrected_rows.append(corrected_series)
                    break
            else:
                corrected_rows.append(row)
        else:
            corrected_rows.append(row)
    
    return pd.DataFrame(corrected_rows, columns=df_backlog.columns)

def extract_text_from_pdf(pdf_bytes):
    """
    Haalt tekst uit een PDF-bestand.
    """
    try:
        with pdfplumber.open(pdf_bytes) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        return text
    except Exception as e:
        pass
        return ""

def extract_text_from_excel(excel_bytes):
    """
    Haalt tekst uit een Excel (.xlsx) bestand.
    """
    try:
        with BytesIO(excel_bytes) as buffer:
            xls = pd.ExcelFile(buffer)
            extracted_text = []
            
            for sheet_name in xls.sheet_names:  # Loop door alle werkbladen
                df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str)  # Lees als tekst
                text = df.applymap(str).values.flatten()  # Converteer alle cellen naar strings
                extracted_text.append(f"--- Blad: {sheet_name} ---\n" + "\n".join(text))

            return "\n".join(extracted_text) if extracted_text else ""

    except Exception as e:
        st.error(f"Fout bij tekstextractie uit Excel: {e}")
        return ""

def extract_text_from_docx(docx_bytes):
    """
    Haalt tekst uit een Word (.docx) bestand.
    """
    try:
        with BytesIO(docx_bytes) as buffer:
            doc = Document(buffer)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except Exception as e:
        st.error(f"Fout bij tekstextractie uit Word: {e}")
        return ""

def extract_text_from_rtf(rtf_bytes):
    """
    Haalt tekst uit een RTF-bestand.
    """
    try:
        with BytesIO(rtf_bytes) as buffer:
            rtf_content = buffer.read().decode("utf-8", errors="ignore")
            text = rtf_to_text(rtf_content)
        return text
    except Exception as e:
        st.error(f"Fout bij tekstextractie uit RTF: {e}")
        return ""

def extract_text_from_doc(doc_bytes):
    """
    Haalt tekst uit een .doc-bestand door tijdelijk op te slaan.
    """
    try:
        with NamedTemporaryFile(delete=True, suffix=".doc") as tmp:
            tmp.write(doc_bytes)
            tmp.flush()
            text = textract.process(tmp.name).decode("utf-8", errors="ignore")
        return text
    except Exception as e:
        st.error(f"Fout bij tekstextractie uit DOC: {e}")
        return ""

def extract_text_from_xls(xls_bytes):
    """
    Haalt tekst of waarden uit een .xls-bestand.
    """
    try:
        with BytesIO(xls_bytes) as buffer:
            workbook = xlrd.open_workbook(file_contents=buffer.read())
            result = []
            for sheet in workbook.sheets():
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    if any(cell != '' for cell in row):
                        result.append(" | ".join(str(cell) for cell in row))
        return "\n".join(result)
    except Exception as e:
        st.error(f"Fout bij tekstextractie uit XLS: {e}")
        return ""



def extract_pdf_to_dataframe(pdf_reader, use_gpt_extraction):
    try:
        # **Stap 1: Controleer of er een tabel in de PDF staat**
        table_found = False  # Flag om bij te houden of een tabel is gevonden
        first_table = None  # Variabele om eerste gevonden tabel op te slaan

        with pdfplumber.open(pdf_reader) as pdf:
            for i, page in enumerate(pdf.pages):
                table = page.extract_table()
                if table:
                    table_found = True  # Markeer dat er een tabel is gevonden
                    first_table = table  # Sla de eerste gevonden tabel op
                    break  # Eén tabel is genoeg om de check te voltooien
        
        # **Toon de uitkomst in de UI**
        if table_found and first_table:
            st.success("✅ Een tabel is gevonden in de PDF.")
            df_table = pd.DataFrame(first_table[1:], columns=first_table[0])  # Eerste rij als header gebruiken
            
            # **Debugging Stap**: Controleer of er duplicate indexwaarden zijn

            if df_table.index.duplicated().any():
                st.error("⚠ Waarschuwing: Dubbele indexen gedetecteerd in de tabel!")
                df_table = df_table.reset_index(drop=True)  # Fix index probleem
            
            st.dataframe(df_table)  # Toon de tabel in de UI
            return df_table  # Return de tabel als dataframe
        else:
            if not table_found:
                st.warning("✨ Geen tabel gevonden.")
            
                if use_gpt_extraction:
                    progress_bar = st.progress(0)  # Start een lege progress bar
                    
                    for percent_complete in range(0, 101, 10):  # Laat de balk oplopen van 0% naar 100%
                        time.sleep(0.5)  # Wacht 0.5 seconden per stap (kan worden aangepast)
                        progress_bar.progress(percent_complete)
                    
                    # Voer nu de AI-extractie uit
                    document_text = extract_text_from_pdf(pdf_reader)
                    relevant_data = extract_data_with_gpt(document_text)
                    
                    # Verwijder de progress bar en geef succesmelding
                    progress_bar.empty()
                    st.success("✅ AI-extractie voltooid!")
                    st.code(relevant_data, language="json")
                    
                    # **Controleer of de respons een geldige DataFrame is**
                    if isinstance(relevant_data, pd.DataFrame) and not relevant_data.empty:
                        st.success("✅ AI-extractie voltooid!")

                        st.dataframe(relevant_data)
                        return relevant_data  # Direct GPT-resultaat retourneren
                    else:
                        st.error("❌ Fout bij GPT-extractie: De gegenereerde data is niet geldig.")
                        return pd.DataFrame()  # Voorkom crashes door een lege DataFrame terug te geven
                    
                    st.success("✅ AI-extractie voltooid!")

                    st.dataframe(relevant_data)
                    return relevant_data  # Direct GPT-resultaat retourneren

                else:
                    st.warning("⚠ Geen tabel gevonden, en AI-extractie is niet ingeschakeld.")


            # **Fallback naar tekstextractie als er geen tabel is gevonden**
            with pdfplumber.open(pdf_reader) as pdf:
                lines = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        lines.extend(text.split("\n"))
    
            structured_data = []
            current_category = "0"  # Fallback waarde als er geen categorie is
            category_pattern = re.compile(r"(\d{1,2}-\s*\d{1,2}[A-Z]?-\s*\w+)|(\d{1,2}[A-Z]?\s*-\s*\w+)")

            for line in lines:
                line = line.strip()
                if category_pattern.match(line):
                    current_category = line.replace(":", "")
                    continue
    
                if re.search(r"\bTotaal:?\b", line, re.IGNORECASE):
                    continue
    
                columns = re.split(r'\s+', line)
                if len(columns) >= 5:
                    structured_data.append([current_category] + columns)
    
            if structured_data:
                max_columns = max(len(row) for row in structured_data)
                column_names = ["Categorie"] + [f"Kolom_{i}" for i in range(1, max_columns)]
                structured_data = [row + [""] * (max_columns - len(row)) for row in structured_data]
                df = pd.DataFrame(structured_data, columns=column_names)
    
                header_row = None
                for i in range(min(3, len(df))):  # Maximaal 3 rijen doorzoeken voor headers
                    potential_headers = df.iloc[i].astype(str).str.lower().str.strip()
                    if any(potential_headers.isin([
                        "artikelnaam", "artikel", "product", "type", "article", "samenstelling",
                        "hoogte", "height", "h",
                        "breedte", "width", "b",
                        "aantal", "quantity", "qty", "stuks"
                    ])):
                        header_row = i
                        break
    
                if header_row is not None:
                    # Controleer en hernoem dubbele kolomnamen
                    df.columns = df.iloc[header_row].astype(str).str.strip()  # Strip spaties en zet om naar string
                    df = df.drop(df.index[:header_row + 1])  # Verwijder header-rij
                    
                    # Hernoem dubbele kolommen
                    if df.columns.duplicated().any():
                        st.error(f"⚠ Fout: Dubbele kolomnamen gevonden: {df.columns[df.columns.duplicated()].tolist()}")
                    
                        new_columns = []
                        col_count = {}
                        for col in df.columns:
                            if col in col_count:
                                col_count[col] += 1
                                new_columns.append(f"{col}_{col_count[col]}")  # Voeg index toe aan dubbele kolommen
                            else:
                                col_count[col] = 1
                                new_columns.append(col)
                    
                        df.columns = new_columns  # Update kolomnamen
                        st.success("✅ Dubbele kolomnamen hernoemd.")

                    df = df.drop(df.index[:header_row + 1])
                else:
                    st.warning("⚠ Geen header herkend, eerste rij als header gebruikt.")
                    df.columns = df.iloc[0]
                    df = df.drop(df.index[0])
                
                # **Debugging Stap**: Controleer of de index uniek is



                if not df.index.is_unique:
                    st.error("⚠ Waarschuwing: Niet-unieke indexwaarden gevonden vóór reset. Fix index...")
                    st.write("Dubbele indexen gevonden:", df.index[df.index.duplicated()].tolist())  # Debug info
                    df = df.loc[~df.index.duplicated(keep='first')].reset_index(drop=True)
                else:
                    df = df.reset_index(drop=True)

                # Debugging: Controleer opnieuw na reset
                if not df.index.is_unique:
                    st.error("⚠ Probleem na reset: Index is nog steeds niet uniek!")
                    st.write("Huidige indexstatus:", df.index)

                # Extra: Print de kolommen en rijen om te checken of data correct is


                df.columns = df.columns.str.lower()
    
                if "aantal" not in df.columns:
                    st.error("⚠ Kolom 'aantal' niet gevonden in de PDF.")
                    st.write("Herkende kolommen:", df.columns.tolist())
                    return pd.DataFrame()
    
                if df.shape[0] > 2:
                    df = pd.concat([
                        df.iloc[:2],  
                        df.iloc[2:][~df.iloc[2:].apply(
                            lambda row: row.astype(str).str.contains(r"\b(Aantal|Breedte|Hoogte)\b", case=False).any(), axis=1)
                        ]
                    ]).reset_index(drop=True)
    
                for col in ["aantal", "breedte", "hoogte"]:
                    if col in df.columns:
                        df[col] = pd.to_numeric(df[col], errors="coerce")  
                
                if "df_current" not in st.session_state:
                    st.session_state.df_current = df.copy()
                if "batch_number" not in st.session_state:
                    st.session_state.batch_number = 1
                if "next_df" not in st.session_state:
                    st.session_state.next_df = None
                
                if st.session_state.next_df is not None:
                    st.session_state.df_current = st.session_state.next_df.copy()
                    st.session_state.next_df = None  
                
                df_current = st.session_state.df_current
                
                df_backlog = df_current[
                    df_current["aantal"].isna() | (df_current["aantal"] <= 0) |
                    df_current["breedte"].isna() | (df_current["breedte"] < 100) |
                    df_current["hoogte"].isna() | (df_current["hoogte"] < 100)
                ]
                
                if not df_backlog.empty:
                    df_corrected = correct_backlog_rows(df_current)
                    df_current.update(df_corrected)
                
                df_bulk = df_current.loc[
                    ~df_current.index.isin(df_backlog.index)
                ].copy()
    
                st.write("✅ **Verwerkte gegevens:**")
                st.dataframe(df_corrected)
                
                return df_corrected  
    
            else:
                st.warning("Geen gegevens gevonden in de PDF om te verwerken test.")
    
    except Exception as e:
        st.error(f"Fout bij het extraheren van PDF-gegevens: {e}")
        return pd.DataFrame()






        
def extract_latest_email(body):
    """
    Extracts only the latest email from an email thread.
    It detects the start of a new email using the pattern 'Van:' followed by 'Verzonden:'.
    """
    email_parts = re.split(r'Van:.*?Verzonden:.*?Aan:.*?Onderwerp:', body, flags=re.DOTALL)
    if email_parts:
        latest_email = email_parts[0].strip()
        return latest_email
    else:
        return body.strip()



def debug_check_tables(doc_bytes):
    """ Controleert of er tabellen in het DOCX-bestand zijn en toont extra statistieken. """
    doc = Document(BytesIO(doc_bytes))
    num_tables = len(doc.tables)

    print(f"📊 Aantal tabellen in het DOCX-bestand: {num_tables}")

    total_table_cells = 0
    total_text_lines = 0
    total_words = 0
    total_chars = 0

    # **Tabelinformatie**
    for i, table in enumerate(doc.tables):
        st.write(f"📂 Tabel {i+1}:")
        for row in table.rows:
            cell_values = [cell.text.strip() for cell in row.cells]
            st.write(cell_values)
            total_table_cells += len(cell_values)
        st.write("="*50)  # Visuele scheiding tussen tabellen

    if num_tables == 0:
        st.write("❌ Geen tabellen gevonden in het document!")

    # **Tekstinformatie**
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            total_text_lines += 1
            total_chars += len(text)
            total_words += len(text.split())

    # **Toon de statistieken**
    st.write(f"📃 Aantal regels tekst: {total_text_lines}")
    st.write(f"🔤 Aantal tekens: {total_chars}")
    st.write(f"📝 Aantal woorden: {total_words}")

    if num_tables > 0:
        st.write(f"📦 Totaal aantal cellen in tabellen: {total_table_cells}")

def convert_docx_to_xlsx(doc_bytes):
    """
    Converteer een DOCX-bestand naar een Excel-bestand en neem ALLE inhoud mee.
    """
    # **Voer eerst de debug-check uit**
    debug_check_tables(doc_bytes)

    # Maak een tijdelijk Excel-bestand aan
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
        excel_output_path = temp_file.name

    with pd.ExcelWriter(excel_output_path, engine="openpyxl") as writer:
        table_count = 0
        has_data = False

        # **Stap 1: Controleer tabellen**
        if len(doc.tables) > 0:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):  # Lege rijen negeren
                        rows.append(cells)

                if rows:
                    df = pd.DataFrame(rows)
                    df.columns = [f"Kolom_{i+1}" for i in range(df.shape[1])]  # Fallback headers

                    table_count += 1
                    df.to_excel(writer, sheet_name=f"Tabel_{table_count}", index=False)
                    has_data = True
        else:
            st.write("❌ Geen tabellen gevonden! We proberen tekst als tabel te verwerken.")

        # **Stap 2: Als er geen tabellen zijn, probeer tekstregels als tabel te extraheren**
        structured_data = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Probeer regels te splitsen op tab of meerdere spaties (lijkt op een tabel)
                row = [t.strip() for t in text.split("\t")]  # Eerst op tabs splitsen
                if len(row) < 2:  # Als te weinig kolommen, probeer spaties
                    row = [t.strip() for t in text.split("  ")]  # Twee of meer spaties als scheiding
                structured_data.append(row)

        # **Stap 3: Sla gestructureerde tekst op als tabel**
        if structured_data:
            df_text = pd.DataFrame(structured_data)
            df_text.columns = [f"Kolom_{i+1}" for i in range(df_text.shape[1])]  # Fallback headers
            df_text.to_excel(writer, sheet_name="Gestructureerde Tekst", index=False)
            has_data = True

        # **Stap 4: Voeg een minimale zichtbare sheet toe als er geen gegevens zijn**
        if not has_data:
            df_empty = pd.DataFrame({"Melding": ["Geen data gevonden"]})
            df_empty.to_excel(writer, sheet_name="Leeg Document", index=False)

    return excel_output_path




def extract_data_with_gpt(prompt):
    """
    Verstuurt een tekstprompt naar GPT en retourneert een geformatteerde tekstoutput per offerteregel.
    """
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": (
                    "Je bent een geavanceerde extractietool die glassamenstellingen uit een bestekformulier extraheert en deze omzet naar een tekstuele lijst.\n"
                    "Formatteer de output volgens het volgende formaat:\n"
                    "[aantal]x {[omschrijving]} [breedte]x[hoogte] [Warmedge] \n"
                    "Verwijder eventuele spaties uit de omschrijving, voorbeelden van omschrijvingen zijn '4-15-4', '8-18A-44.2', '33/1-33/1' of '5-5'.\n"
                    "Warmedge betekent een zwarte, of warmedge spacer/afstandhouder bij isolatieglas. Je kunt het herkennen aan Warm-E, warmedge, zwarte spacer, zwarte afstandhouder etc. Het is niet standaard of basic. Als van toepassing, eindig de regel met 'WE'. Geen warmedge, plaats dan geen extra tekst.\n"
                    "Mocht een regel geen omschrijving hebben, neem je de omschrijving van de voorgaande regel.\n"
                    "Het volgende is enorm belangrijk: Wanneer de productgroep genoemd wordt, bijv. 'Eclaz One', 'SS Zero', 'Alfa', 'Zonwerend' 'SKN', plaats deze dan bóven de regel tussen **, bijv. *Eclaz One*. Alle volgende regels zullen dan onder die productgroep vallen.\n"
                    "De productgroep kan na het artikelnummer worden genoemd, bovenaan als kop, of iets verderop tussen kenmerken van de samenstelling: bijv. 6#-15-4 ZTA: 40%  U:1.0 W/m2k | Energy | Waarbij energy dan de productgroep is.\n"
                    "Houd je strikt aan dit formaat zonder extra uitleg, JSON, Markdown of aanvullende tekst."
                )},
                {"role": "user", "content": prompt}
            ]
        )
        
        extracted_text = response.choices[0].message.content.strip()


        # **Stap 1: Opschonen van extra tekens**
        extracted_text = extracted_text.replace("```plaintext", "").replace("```", "").strip()

        # **Stap 2: Opslaan en weergeven**
        st.session_state["geformatteerde_output"] = extracted_text  # ✅ Sla de geformatteerde output op
        st.session_state["customer_input"] = extracted_text  # ✅ Vul customer_input automatisch in
        return extracted_text

    except Exception as e:
        st.error(f"❌ Fout bij GPT-extractie: {e}")
        return ""

def process_single_attachment(selected_name, selected_data):
    ext = Path(selected_name).suffix.lower()
    
    try:
        if ext == ".pdf":
            document_text = extract_text_from_pdf(BytesIO(selected_data))
        elif ext == ".xlsx":
            document_text = extract_text_from_excel(selected_data)
        elif ext == ".docx":
            document_text = extract_text_from_docx(selected_data)
        elif ext == ".rtf":
            document_text = extract_text_from_rtf(selected_data)     
        elif ext == ".doc":
            document_text = extract_text_from_doc(selected_data)  
        elif ext == ".xls":
            document_text = extract_text_from_xls(selected_data)  
        elif ext == ".msg":
            return None
        else:
            st.error(f"Onbekend bestandstype: {ext}. Alleen .pdf, .xls(x), .rtf, en .doc(x) worden ondersteund.")
            return None

        if not document_text:
            st.warning("Kon geen tekst extraheren uit dit bestand.")
            return None

        extracted_data = extract_data_with_gpt(document_text)

        if isinstance(extracted_data, pd.DataFrame) and not extracted_data.empty:
            st.success(f"Data succesvol geëxtraheerd uit {selected_name}:")
            st.dataframe(extracted_data)

            if st.button("📄 Verwerk gegevens naar offerte"):
                handle_mapped_data_to_offer(extracted_data)

            return extracted_data
        else:
            return None

    except Exception as e:
        st.error(f"Fout bij het verwerken van de bijlage: {e}")
        return None


def process_attachment(attachments):
    """
    Verwerkt een bestand of bijlage via verplichte AI-extractie.
    - Bij .msg: kies uit de bijlagen (alleen pdf/xlsx/docx).
    - Bij andere bestanden: verwerk direct.
    """
    valid_extensions = [".pdf", ".xlsx", ".docx", ".msg", ".rtf", ".doc", ".xls"]
    excluded_extensions = (".png", ".jpg", ".jpeg")

    if isinstance(attachments, list):
        # Filter alleen bruikbare bijlagen (geen afbeeldingen)
        valid_attachments = {
            att.longFilename or att.shortFilename: att.data
            for att in attachments
            if (att.longFilename or att.shortFilename).lower().endswith(tuple(valid_extensions))
        }

        if not valid_attachments:
            st.info("Geen geschikte bijlagen gevonden voor extractie (.pdf, .xlsx, .docx). Alleen de e-mailinhoud wordt getoond.")
            return None

        # Dropdown tonen als er meerdere geldige bijlagen zijn
        selected_name = st.selectbox("Kies een bijlage voor AI-extractie:", options=list(valid_attachments.keys()))
        selected_data = valid_attachments[selected_name]

        # Knop om extractie uit te voeren
        if st.button(f"Verwerk bijlage ✨"):
            return process_single_attachment(selected_name, selected_data)

    else:
        # Enkel bestand (geen .msg)
        selected_name = attachments.name
        selected_data = attachments.read()
        ext = Path(selected_name).suffix.lower()

        if ext in excluded_extensions:
            return None  # Afbeeldingen negeren

        if ext not in valid_extensions:
            st.warning("Alleen .pdf, .xls(x), .rtf of .docx bestanden kunnen verwerkt worden.")
            return None

        # Direct verwerken zonder dropdown
        return process_single_attachment(selected_name, selected_data)

    return None



st.sidebar.markdown("---")  # Scheidingslijn voor duidelijkheid  

# File uploader alleen beschikbaar in de uitklapbare invoeropties
with st.sidebar.expander("Upload document", expanded=True):
    uploaded_file = st.file_uploader(
        "Upload een Outlook (.msg), PDF, Word of Excel bestand", 
        type=["msg", "pdf", "xlsx", "docx", "rtf", "xls"]
    )

    if uploaded_file:
        ext = Path(uploaded_file.name).suffix.lower()

        if ext == ".msg":
            with open("uploaded_email.msg", "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            try:
                msg = extract_msg.Message("uploaded_email.msg")
                msg_subject = msg.subject
                msg_sender = msg.sender
                full_email_body = msg.body
                latest_email = extract_latest_email(full_email_body)
                msg_body = latest_email
                email_body = msg_body

                # Zet onderwerp automatisch als klantreferentie
                if msg_subject:
                    try:
                        if not st.session_state.get("customer_reference") or not customer_reference.strip():
                            clean_subject = re.sub(r"^(FW:|RE:)\s*", "", msg_subject.strip(), flags=re.IGNORECASE)
                            st.session_state["customer_reference"] = clean_subject
                            st.rerun()
                    except Exception as e:
                        st.error(f"Fout bij het verwerken van de klantreferentie: {e}")

                st.subheader("Berichtinformatie")
                st.write(f"**Onderwerp:** {msg_subject}")
                st.write(f"**Afzender:** {msg_sender}")
                st.write("**Inhoud van het bericht:**")
                st.text(msg_body)

                st.subheader("Bijlagen:")
                if msg.attachments:
                    relevant_data = process_attachment(msg.attachments)
                else:
                    st.info("Geen bijlagen gevonden. Verwerk de tekst in de mail met BullsAI knop")

            except Exception as e:
                st.error(f"Fout bij het verwerken van het bestand: {e}")
        
        else:
            # Los bestand (geen .msg), direct verwerken
            st.subheader("Bestandverwerking")
            extracted_data = process_attachment(uploaded_file)
            if extracted_data is None:
                st.info("Kon geen gegevens extraheren uit dit bestand.")
    else:
        st.info("Upload een bestand om te beginnen.")


# Gebruikersinvoer (wordt automatisch ingevuld vanuit "Geformatteerde output")
customer_input = st.sidebar.text_area(
    "Voer hier handmatig het klantverzoek in.", 
    value=st.session_state.get("customer_input", ""),  # ✅ Haalt de waarde op als die er is
    height=200
)


# Functie om tekstinvoer te verwerken
def handle_text_input(input_text):
    matched_articles = [(term, synonym_dict[term]) for term in synonym_dict if term in input_text]

    if matched_articles:
        response_text = "Bedoelt u de volgende samenstellingen:"
        for term, article_number in matched_articles:
            description, _, _, _, _ = find_article_details(lookup_article_number,current_productgroup=current_productgroup)
            if description:
                response_text += f"- {description} met artikelnummer {article_number}\n"

        response_text += "?"
        st.sidebar.write(response_text)
    else:
        st.sidebar.warning("Geen gerelateerde artikelen gevonden. Gelieve meer details te geven.")

# Functie om offerte als PDF te genereren
def generate_pdf(df):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from io import BytesIO

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []

    # Styles
    styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        'HeaderStyle', parent=styles['Heading1'], fontSize=17, alignment=TA_CENTER, textColor=colors.black
    )
    normal_style = ParagraphStyle(
        'NormalStyle', parent=styles['Normal'], fontSize=11, alignment=TA_LEFT, textColor=colors.black
    )
    right_aligned_style = ParagraphStyle(
        'RightAlignedStyle', parent=styles['Normal'], fontSize=11, alignment=TA_LEFT, textColor=colors.black
    )

    # Header
    elements.append(Paragraph("Vandaglas - Offerte", header_style))
    elements.append(Spacer(1, 12))

    # Introductietekst
    elements.append(Paragraph(
        "Beste klant,<br/><br/>"
        "Hartelijk dank voor uw prijsaanvraag. Hieronder vindt u onze offerte. Wij hopen u een passend aanbod te hebben gedaan. "
        "Uw contactpersoon, Job, geeft graag nog een toelichting en beantwoordt eventuele vragen.<br/><br/>"
        "Met vriendelijke groet,<br/>"
        "Vandaglas",
        normal_style
    ))
    elements.append(Spacer(1, 24))

    # Tabel header
    data = [["Artikelnaam", "Breedte", "Hoogte", "Aantal", "Prijs p/s", "M2 p/s", "Totaal M2", "Totaal"]]

    # Voeg gegevens uit df toe aan tabel
    for index, row in df.iterrows():
        if all(col in row for col in ['Artikelnaam', 'Breedte', 'Hoogte', 'Aantal', 'RSP', 'M2 p/s', 'M2 totaal']):
            data.append([
    row['Artikelnaam'],
    row['Breedte'],
    row['Hoogte'],
    row['Aantal'],
    row['Prijs_backend'],
    f"{float(str(row['M2 p/s']).replace('m²', '').replace(',', '.').strip()):.2f} m2" if pd.notna(row['M2 p/s']) else None,
    f"{float(str(row['M2 totaal']).replace('m²', '').replace(',', '.').strip()):.2f} m2" if pd.notna(row['M2 totaal']) else None,
    f"{round(float(str(row['Prijs_backend']).replace('€', '').replace(',', '.').strip()) * float(row['Aantal']) * float(str(row['M2 p/s']).replace('m²', '').replace(',', '.').strip()), 2):,.2f}" if pd.notna(row['Prijs_backend']) and pd.notna(row['Aantal']) else None
])


    # Maak de tabel
    table = Table(data, repeatRows=1, colWidths=[150, 45, 45, 45, 45, 45, 45, 60])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 24))

    # Eindtotaal, BTW, Te betalen
    total_price = df.apply(lambda row: round(float(str(row['Prijs_backend']).replace('€', '').replace(',', '.').strip()) * float(str(row['M2 totaal']).replace('m²', '').replace(',', '.').strip()), 2) if pd.notna(row['Prijs_backend']) and pd.notna(row['M2 totaal']) else 0, axis=1).sum()
    btw = total_price * 0.21
    te_betalen = total_price + btw

    # Maak klein tabelletje voor totalen
    totals_data = [
        ["Eindtotaal:", f"€ {total_price:.2f}"],
        ["BTW (21%):", f"€ {btw:.2f}"],
        ["Te betalen:", f"€ {te_betalen:.2f}"]
    ]
    totals_table = Table(totals_data, colWidths=[100, 100])
    totals_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))

    from reportlab.lib.units import inch
    elements.append(Spacer(1, 0.5 * inch))
    totals_table = Table(totals_data, colWidths=[100, 100], hAlign='RIGHT')
    totals_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    elements.append(Spacer(1, 3 * inch))
    elements.append(totals_table)
  

    # Bouwelementen aan document
    doc.build(elements)
    buffer.seek(0)
    return buffer

# # Offerte Genereren tab
# with tab1:
#     # Knop om GPT-chat te versturen
#     if st.sidebar.button("Vertaal chat naar offerte"):
#         try:
#             handle_gpt_chat()
#         except Exception as e:
#             st.sidebar.error(f"Er is een fout opgetreden: {e}")

#     # Knop om de e-mail te vertalen naar een offerte
#     if st.sidebar.button("Vertaal mail naar offerte"):
#         try:
#             handle_email_to_offer(email_body)
#         except Exception as e:
#             st.error(f"Fout bij het verwerken van de e-mail: {e}")

# Offerte Genereren tab
with tab1:
    # Een container voor de knop en de informatiebalk
    with st.sidebar.container():   
        relevant_data = None

        # Verwerk de bijlage zodra deze is geüpload
        if uploaded_file is not None:
            attachment_name = uploaded_file.name
            relevant_data = process_attachment(uploaded_file)
        
        # Eén knop om de acties uit te voeren
        if st.sidebar.button("BullsAI 🚀"):
            actie_uitgevoerd = False
        
            # Spinner toevoegen rond alle acties
            with st.spinner("BullsAI is bezig met de verwerking..."):
                # Probeer de eerste actie (tekstvak naar offerte)
                try:
                    handle_gpt_chat()
                    actie_uitgevoerd = True
                except Exception:
                    pass  # Fout negeren en doorgaan naar de volgende actie
        
                # Als de eerste actie niet slaagt, probeer de tweede (bijlage mail)
                if not actie_uitgevoerd and relevant_data is not None:
                    try:
                        handle_mapped_data_to_offer(relevant_data)
                        actie_uitgevoerd = True
                    except Exception:
                        pass  # Fout negeren en doorgaan naar de volgende actie
        
                # Als de tweede actie niet slaagt, probeer de derde (mail naar offerte)
                if not actie_uitgevoerd:
                    try:
                        handle_email_to_offer(email_body)
                        actie_uitgevoerd = True
                    except Exception:
                        pass  # Fout negeren
        
            # Eindstatus bepalen
            if actie_uitgevoerd:
                pass
            else:
                st.error("BullsAI heeft geen gegevens kunnen verwerken.")

# Voeg rijnummers toe aan de offerte DataFrame als deze nog niet bestaat
if 'Rijnummer' not in st.session_state.offer_df.columns:
    st.session_state.offer_df.insert(0, 'Rijnummer', range(1, len(st.session_state.offer_df) + 1))

    


# Offerte Genereren tab
with tab1:    
    with col6:
        # Voeg een knop toe om de offerte als PDF te downloaden
        if totaal_bedrag > 25000:
            st.button("Download offerte als PDF", key='download_pdf_button', disabled=True)
            st.button("Autoriseer offerte", key='authorize_offer_button')
        else:
            if st.button("Download offerte als PDF", key='download_pdf_button'):
                pdf_buffer = generate_pdf(st.session_state.offer_df)
                st.download_button(label="Download PDF", data=pdf_buffer, file_name="offerte.pdf", mime="application/pdf")
    
        # Knop om offerte op te slaan in database
        if st.button("Sla offerte op"):
            try:
                # Haal de ingelogde Windows-gebruikersnaam op
                import os
                windows_user = getpass.getuser() if getpass.getuser() else "Onbekende gebruiker"
        
                # Zoek het hoogste offertenummer
                if not st.session_state.saved_offers.empty:
                    max_offer_number = st.session_state.saved_offers['Offertenummer'].max()
                    offer_number = max_offer_number + 1
                else:
                    offer_number = 1
        
                # Bereken eindtotaal
                if all(col in st.session_state.offer_df.columns for col in ['RSP', 'M2 totaal']):
                    eindtotaal = st.session_state.offer_df.apply(
                        lambda row: float(row['RSP']) * float(row['M2 totaal']) if pd.notna(row['RSP']) and pd.notna(row['M2 totaal']) else 0,
                        axis=1
                    ).sum()
                else:
                    eindtotaal = 0
        
                # Voeg offerte-informatie toe aan opgeslagen offertes in sessie
                offer_summary = pd.DataFrame({
                    'Offertenummer': [offer_number],
                    'Klantnummer': [str(st.session_state.customer_number)],
                    'Eindbedrag': [eindtotaal],
                    'Datum': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                    'Gebruiker': [windows_user]  # Voeg gebruikersnaam toe
                })
                st.session_state.saved_offers = pd.concat([st.session_state.saved_offers, offer_summary], ignore_index=True)
        
                # Voeg offertenummer en gebruikersnaam toe aan elke regel in de offerte
                st.session_state.offer_df['Offertenummer'] = offer_number
                st.session_state.offer_df['Gebruiker'] = windows_user
        
                # Opslaan in database
                conn = create_connection()
        #        cursor = conn.cursor()
        
                try:
                    # Voeg elke rij van de offerte toe aan de database
                    for index, row in st.session_state.offer_df.iterrows():
                        cursor.execute("""
                        INSERT INTO Offertes (Offertenummer, Rijnummer, Artikelnaam, Artikelnummer, Spacer, Breedte, Hoogte, Aantal, 
                                              M2_per_stuk, M2_totaal, RSP, SAP_Prijs, Handmatige_Prijs, Min_prijs, Max_prijs, 
                                              Prijs_backend, Verkoopprijs, Source, Datum, Gebruiker)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """, (
                            row['Offertenummer'], index + 1, row['Artikelnaam'], row['Artikelnummer'], row['Spacer'],
                            row['Breedte'], row['Hoogte'], row['Aantal'], row['M2 p/s'], row['M2 totaal'], 
                            row['RSP'], row['SAP Prijs'], row['Handmatige Prijs'], row['Min_prijs'], row['Max_prijs'], 
                            row['Prijs_backend'], row['Verkoopprijs'], row['Source'], datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            windows_user  # Sla gebruikersnaam op
                        ))
        
                    conn.commit()
                    st.success(f"Offerte {offer_number} succesvol opgeslagen door {windows_user}.")
                except sqlite3.Error as e:
                    st.error(f"Fout bij het opslaan in de database: {e}")
                finally:
                    conn.close()
            except Exception as e:
                st.error(f"Er is een fout opgetreden: {e}")


if 'edited_df' in locals() and not edited_df.equals(st.session_state.offer_df):
    edited_df = edited_df.copy()
    edited_df = update_offer_data(edited_df)
    st.session_state.offer_df = edited_df


with tab1:
    # Checkbox voor het creëren van een Opportunity
    if st.checkbox("Creeer Opportunity"):
        # Indeling in vier kolommen
        col1, col2, col3, col4 = st.columns(4)
    
        # Velden in de eerste kolom
        with col1:
            name = st.text_input("Name (gevuld met customer_reference):", value=customer_reference)
            account_id = st.text_input("AccountID:", value="001KI0000084Q8VYAU")
            stage_name = st.selectbox(
                "StageName:",
                options=["RFQ / Initial Quote", "Customer is fixed", "Negotiation", "Verbal Agreement"],
                index=0,  # Standaard geselecteerde waarde
            )
            close_date = st.date_input(
                "CloseDate (datum vandaag + 2 weken):",
                value=date.today() + timedelta(weeks=2),
            )
            amount = st.number_input("Amount (gevuld met totaal_bedrag):", value=totaal_bedrag)
            description = st.text_area("Description (gevuld met customer_reference):", value=customer_reference)
    
        # Knop om de Opportunity aan te maken
        if st.button("Opportunity aanmaken"):
            try:
                # Opportunity-gegevens
                opportunity_data = {
                    "Name": name,
                    "AccountId": account_id,
                    "StageName": stage_name,
                    "CloseDate": close_date.isoformat(),
                    "Amount": amount,
                    "Description": description,
                }
    
                # Opportunity aanmaken in Salesforce
                resultaat = sf.Opportunity.create(opportunity_data)
    
                # Maak een hyperlink naar de Salesforce Opportunity
                opportunity_id = resultaat['id']
                salesforce_url = f"https://vandaglasnl--qa.sandbox.my.salesforce.com/{opportunity_id}"
                hyperlink = f"[{customer_reference}]({salesforce_url})"
    
                st.success(f"Opportunity succesvol aangemaakt! {hyperlink}")
            except Exception as e:
                st.error(f"Fout bij het aanmaken van de Opportunity: {e}")


# Opgeslagen Offertes tab


with tab2:

    # Twee kolommen maken
    col1, col2 = st.columns([1, 1])  
    with col1:
        # Expander onder AgGrid met een gefilterde DataFrame-weergave
        with st.expander("⚡ SAP format", expanded=True):
            if "offer_df" in st.session_state:
                filtered_df = st.session_state.offer_df[["Artikelnummer", "Aantal", "Breedte", "Hoogte", "Spacer"]].copy()

                # Bewaar originele Spacer-kolom voor logica
                original_spacer = filtered_df["Spacer"]

                # Haal alleen het numerieke deel uit Spacer
                filtered_df["Spacer"] = filtered_df["Spacer"].str.extract(r'(\d+)')

                # Voeg 3 lege kolommen toe vóór Spacertype
                filtered_df[""] = ""
                filtered_df["  "] = ""
                filtered_df["   "] = ""

                # Voeg kolom Spacertype toe: 13 als "warm edge" in oorspronkelijke waarde
                filtered_df["Spacertype"] = original_spacer.str.contains("warm edge", case=False, na=False).map({True: 13, False: ""})

                # Voeg 3 lege kolommen toe na Spacertype
                filtered_df["    "] = ""
                filtered_df["     "] = ""
                filtered_df["      "] = ""

                # Voeg kolom Spacerkleur toe: "Zwart" als Spacertype 13
                filtered_df["Spacerkleur"] = filtered_df["Spacertype"].apply(lambda x: "2" if x == 13 else "")

                # Toon de gefilterde DataFrame
                st.dataframe(filtered_df, use_container_width=True)

                # Zet DataFrame om naar tab-gescheiden tekst zonder headers
                table_text = filtered_df.to_csv(index=False, sep="\t", header=False).strip()

                # JavaScript-code om de tabel naar het klembord te kopiëren
                copy_js = f"""
                <script>
                    function copyToClipboard() {{
                        let text = `{table_text}`;  // Alleen de data, geen headers
                        navigator.clipboard.writeText(text).then(() => {{
                            alert("✅ Gegevens gekopieerd naar klembord!");
                        }}).catch(err => {{
                            alert("❌ Fout bij kopiëren: " + err);
                        }});
                    }}
                </script>
                <button onclick="copyToClipboard()" style="padding:8px 16px; background:#4CAF50; color:white; border:none; border-radius:4px; cursor:pointer;">
                    📋 Kopieer naar klembord
                </button>
                """

                # Weergeven van de knop via Streamlit's componenten
                st.components.v1.html(copy_js, height=50)
            else:
                st.warning("Geen gegevens beschikbaar om weer te geven.")







with tab3:
    # Layout met twee kolommen
    col1, col2 = st.columns(2)
    
    # Linkerkolom: Tabel met synoniemen beoordelen
    with col1:
        st.markdown("### Beoordeel output AI ✨")
    
        # Controleer of offer_df beschikbaar is in sessiestatus
        if "offer_df" in st.session_state and not st.session_state.offer_df.empty:
            # Filter regels met "Source" = "interpretatie", "niet gevonden" en "GPT"
            interpretatie_rows = st.session_state.offer_df[st.session_state.offer_df["Source"].isin(["GPT", "interpretatie", "niet gevonden"])]

            # Houd alleen unieke rijen op basis van combinatie van kolommen
            interpretatie_rows = interpretatie_rows.drop_duplicates(subset=["Artikelnaam", "Artikelnummer", "fuzzy_match", "original_article_number"])
        else:
            interpretatie_rows = pd.DataFrame()  # Lege DataFrame als fallback
    
        if interpretatie_rows.empty:
            st.info("Er zijn geen AI regels om te beoordelen.")
        else:
            # Maak een kopie van de DataFrame om bewerkingen uit te voeren
            beoordeling_tabel = interpretatie_rows.copy()

            # Pas de voorwaarden toe op de kolommen
            beoordeling_tabel.loc[beoordeling_tabel["Source"] == "niet gevonden", "Artikelnaam"] = ""
            beoordeling_tabel["Bron"] = beoordeling_tabel["Source"].replace({
                "interpretatie": "✨",
                "GPT": "✨",
                "niet gevonden": ""  # Maak "Bron" leeg
            })

            # Selecteer en hernoem de kolommen
            beoordeling_tabel = beoordeling_tabel[["Artikelnaam", "Artikelnummer", "fuzzy_match", "original_article_number", "Bron"]].fillna("")
            beoordeling_tabel.rename(columns={
                "Artikelnaam": "Artikelnaam",
                "Artikelnummer": "Artikelnummer",
                "original_article_number": "Jouw input",
                "fuzzy_match": "Gematcht op",
                "Bron": "Bron"
            }, inplace=True)

            # Configureren van de AgGrid-tabel
            gb = GridOptionsBuilder.from_dataframe(beoordeling_tabel)
            
            # Instellen van een dropdown voor de kolom "Artikelnaam"
            gb.configure_column(
                "Artikelnaam",
                editable=True,
                cellEditor="agRichSelectCellEditor",
                cellEditorParams={
                    "values": list(article_mapping.keys()),  # De mogelijke waarden
                    "searchable": True,  # Laat je typen in de dropdown
                    "suppressKeyboardEvent": False  # Zorgt dat je kunt typen zonder dat de dropdown sluit
                }
            )

            
            # Configureren van de overige kolommen
            gb.configure_column("Artikelnummer", editable=False, hide=True)
            gb.configure_column("Gematcht op", editable=False)
            gb.configure_column("Jouw input", editable=False)
            gb.configure_column("Bron", editable=False)  # Bron wordt leeg als "niet gevonden"

            gb.configure_selection(selection_mode="multiple", use_checkbox=True)
            grid_options = gb.build()

            # Render de AgGrid-tabel
            response = AgGrid(
                beoordeling_tabel,
                gridOptions=grid_options,
                update_mode=GridUpdateMode.MODEL_CHANGED,
                fit_columns_on_grid_load=True,
                theme="material"
            )

    
            # Verwerken van wijzigingen in de tabel
            updated_rows = response["data"]  # Haal de bijgewerkte data op
            for index, row in updated_rows.iterrows():
                # Bijwerken van het artikelnummer op basis van de geselecteerde artikelnaam
                if row["Artikelnaam"] in article_mapping:
                    updated_rows.at[index, "Artikelnummer"] = article_mapping[row["Artikelnaam"]]
    
            # Knop voor accordering
            if st.button("Accordeer synoniem"):
                geselecteerde_rijen = response.get("selected_rows", [])
            
                if isinstance(geselecteerde_rijen, list):  
                    geselecteerde_rijen = pd.DataFrame(geselecteerde_rijen)
            
                if geselecteerde_rijen.empty:
                    st.warning("Geen rijen geselecteerd of response is leeg.")
                else:
                    # Zet alle waarden om naar strings om fouten te voorkomen
                    geselecteerde_rijen = geselecteerde_rijen.astype(str)
            
                    # Maak databaseverbinding
                    engine = create_connection()
                    if engine is None:
                        st.error("Databaseverbinding kon niet worden opgezet. Controleer de instellingen.")
                    else:
                        with engine.connect() as conn:
                            try:
                                # Controleer of de tabel bestaat en maak deze aan indien nodig (met juiste datatypes)
                                query_create_table = text("""
                                    IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'SynoniemenAI')
                                    BEGIN
                                        CREATE TABLE SynoniemenAI (
                                            Synoniem NVARCHAR(255) NOT NULL,
                                            Artikelnummer NVARCHAR(255) NOT NULL,  -- Gewijzigd van TIME naar NVARCHAR
                                            Datum DATE DEFAULT GETDATE(),
                                            PRIMARY KEY (Synoniem, Artikelnummer)
                                        );
                                    END
                                """)
                                conn.execute(query_create_table)
            
                                success_count = 0
                                duplicate_count = 0
            
                                # Loop door de geselecteerde rijen
                                for _, rij in geselecteerde_rijen.iterrows():
                                    synoniem = str(rij.get("Jouw input", "")).strip()
                                    artikelnummer = str(rij.get("Artikelnummer", "")).strip()
                                    datum = pd.Timestamp.now().date()  # Zet naar DATE formaat (YYYY-MM-DD)
            
                                    if synoniem and artikelnummer:
                                        # Controleer of het synoniem al in de database staat
                                        query_check = text("""
                                            SELECT COUNT(*) FROM SynoniemenAI WHERE Synoniem = :synoniem AND Artikelnummer = :artikelnummer;
                                        """)
                                        result = conn.execute(query_check, {"synoniem": synoniem, "artikelnummer": artikelnummer})
                                        exists = result.scalar()
            
                                        if exists == 0:
                                            # Voeg de nieuwe synoniem-toewijzing toe
                                            query_insert = text("""
                                                INSERT INTO SynoniemenAI (Synoniem, Artikelnummer, Datum)
                                                VALUES (:synoniem, :artikelnummer, :datum);
                                            """)
                                            conn.execute(query_insert, {
                                                "synoniem": synoniem,
                                                "artikelnummer": artikelnummer,
                                                "datum": datum  # Correcte DATE waarde
                                            })
                                            success_count += 1
                                        else:
                                            duplicate_count += 1
            
                                conn.commit()
            
                                st.write(f"✅ Succesvol toegevoegd: {success_count}")
                                if duplicate_count > 0:
                                    st.info(f"⚠ {duplicate_count} synoniemen bestonden al en zijn overgeslagen.")
            
                            except Exception as e:
                                st.error(f"Fout bij het opslaan: {e}")


# Rechterkolom: Excel-file uploader in een expander
with col2:
    st.markdown("### Upload synoniemen 🧍‍♂⬌🧍‍♂️")
   
    with st.expander("Upload Synoniemen via Excel ✨"):
        st.markdown("Upload een Excel-bestand met de kolommen: **Artikelnummer** en **Synoniem**.")
    
        uploaded_file = st.file_uploader("Upload een Excel-bestand", type=["xlsx"])
        
        if uploaded_file is not None:
            try:
                # Lees het geüploade Excel-bestand
                df_synoniemen = pd.read_excel(uploaded_file)
    
                # Controleer of het bestand de juiste kolommen heeft
                if "Artikelnummer" in df_synoniemen.columns and "Synoniem" in df_synoniemen.columns:
                    if st.button("Upload🔥"):
                        # Maak verbinding met de database
                        engine = create_connection()
                        success_count = 0
                        error_count = 0
    
                        with engine.connect() as conn:
                            for _, row in df_synoniemen.iterrows():
                                artikelnummer = str(row["Artikelnummer"]).strip()
                                synoniem = str(row["Synoniem"]).strip()
                                datum = pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")  # Huidige tijd
    
                                try:
                                    # Voeg synoniem toe aan de database
                                    query = text("""
                                        INSERT INTO SynoniemenAI (Synoniem, Artikelnummer, Input, Datum, Artikelnaam, Bron) 
                                        SELECT :synoniem, :artikelnummer, :synoniem, :datum, NULL, NULL
                                        WHERE NOT EXISTS (
                                            SELECT 1 FROM SynoniemenAI WHERE Synoniem = :synoniem AND Artikelnummer = :artikelnummer
                                        );
                                    """)
                                    conn.execute(query, {
                                        "synoniem": synoniem,
                                        "artikelnummer": artikelnummer,
                                        "datum": datum
                                    })
                                    success_count += 1
                                except Exception as e:
                                    error_count += 1
                                    st.warning(f"Fout bij uploaden van {artikelnummer}: {e}")
    
                            conn.commit()
    
                        st.write(f"✅ Succesvol toegevoegd: {success_count}")
                        st.write(f"❌ Fouten bij toevoegen: {error_count}")
    
                else:
                    st.error("Het bestand moet de kolommen **'Artikelnummer'** en **'Synoniem'** bevatten.")
            except Exception as e:
                st.error(f"Fout bij het lezen van het bestand: {e}")

with col2:
    def generate_excel():
        """
        Genereer een Excel-bestand met twee tabbladen:
        1. "Synoniemen" met kolommen "Artikelnummer" en "Synoniem"
        2. "Bekende Artikelen" met de volledige artikelenlijst uit Articles.py
        """
        # Data voor tabblad 1
        synonyms_data = pd.DataFrame(columns=["Artikelnummer", "Synoniem"])
    
        # Data voor tabblad 2
        articles_data = pd.DataFrame(article_table)[["Material", "Description"]]   # Zet de geïmporteerde articles-lijst om naar een DataFrame
    
        # Schrijf naar een Excel-bestand met twee tabbladen
        output = BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            synonyms_data.to_excel(writer, sheet_name="Nieuwe synoniemen", index=False)
            articles_data.to_excel(writer, sheet_name="Artikelnummer lijst", index=False)
        
        output.seek(0)
        return output
    
    # Streamlit-interface
    st.markdown("### Download Excel voor synoniemen opvoer ⬇️")        
    
    # Maak het Excel-bestand beschikbaar voor download
    excel_file = generate_excel()
    st.download_button(
        label="Download Excel",
        data=excel_file,
        file_name="Artikelen.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

with tab5:

    # # 🎙️ Spraakherkenning instellen
    # recognizer = sr.Recognizer()
    
    # # Salesforce Connectie
    # SF_USERNAME = os.getenv("SALESFORCE_USERNAME")
    # SF_PASSWORD = os.getenv("SALESFORCE_PASSWORD")
    # SF_SECURITY_TOKEN = os.getenv("SF_SECURITY_TOKEN")
    # SF_DOMAIN = "test"
    
    # # 🎤 WebRTC Instellingen
    # webrtc_ctx = webrtc_streamer(
    #     key="speech-recorder",
    #     mode=WebRtcMode.SENDRECV,
    #     rtc_configuration={"iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]},
    #     media_stream_constraints={"audio": True, "video": False},
    # )
    
    
    # def connect_to_salesforce():
    #     try:
    #         if not SF_USERNAME or not SF_PASSWORD or not SF_SECURITY_TOKEN:
    #             st.error("❌ Salesforce login gegevens ontbreken.")
    #             return None
            
    #         session_id, instance = SalesforceLogin(
    #             username=SF_USERNAME,
    #             password=SF_PASSWORD + SF_SECURITY_TOKEN,
    #             domain=SF_DOMAIN
    #         )
    #         sf = Salesforce(instance=instance, session_id=session_id)
    #         st.success("✅ Salesforce-verbinding geslaagd!")
    #         return sf
    #     except Exception as e:
    #         st.error(f"❌ Salesforce-verbinding mislukt: {e}")
    #         return None
    
    # def fetch_salesforce_accounts_direct(sf):
    #     query = "SELECT Name, ERP_Number__c FROM Account"
    #     response = sf.query_all(query)
    #     return response["records"]
    
    # def save_to_salesforce(sf, account_id, comment):
    #     try:
    #         data = {
    #             "General_comment__c": comment,
    #             "Account__c": account_id
    #         }
    #         sf.Minute_Report__c.create(data)
    #         st.success("✅ Minute report opgeslagen in Salesforce!")
    #     except Exception as e:
    #         st.error(f"❌ Fout bij opslaan in Salesforce: {e}")
    
    # def transcribe_audio(audio_bytes):
    #     try:
    #         audio_file = BytesIO(audio_bytes)
    #         audio_file.name = "audio.wav"
    #         response = openai.Audio.transcribe("whisper-1", audio_file)
    #         return response.get("text", "")
    #     except Exception as e:
    #         st.error(f"❌ Fout bij transcriptie: {e}")
    #         return ""
    
    # # 🚀 UI Begin
    # st.title("🎙️ Spraaknotities opslaan in Salesforce")
    
    # with st.expander("📌 Inspreken en opslaan in Minute Report"):
    
    #     # Salesforce Verbinding
    #     sf = connect_to_salesforce()
    
    #     if sf:
    #         accounts = fetch_salesforce_accounts_direct(sf)
    #     else:
    #         accounts = []
    
    #     if accounts:
    #         accounts_df = pd.DataFrame(accounts).drop(columns="attributes", errors="ignore")
    #         accounts_df.rename(columns={"Name": "Klantnaam", "ERP_Number__c": "Klantnummer"}, inplace=True)
    #         accounts_df["Klantinfo"] = accounts_df["Klantnummer"] + " - " + accounts_df["Klantnaam"]
    #     else:
    #         accounts_df = pd.DataFrame(columns=["Klantnaam", "Klantnummer", "Klantinfo"])
    
    #     selected_account = st.selectbox("Selecteer een account:", accounts_df["Klantinfo"] if not accounts_df.empty else [])
    
    #     # 🎙️ Opname Starten
    #     st.write("🎤 Klik op 'Start' en spreek in:")
    
    #     if webrtc_ctx.audio_receiver:
    #         audio_frames = webrtc_ctx.audio_receiver.get_frames(timeout=1)
            
    #         if audio_frames:
    #             audio = av.AudioFrame.from_ndarray(audio_frames[0].to_ndarray(), format="s16")
    #             audio_bytes = audio.to_ndarray().tobytes()
    
    #             # Transcriptie uitvoeren
    #             transcribed_text = transcribe_audio(audio_bytes)
    #             st.text_area("📝 Getranscribeerde tekst:", transcribed_text, height=150)
    
    #             # 💾 Opslaan in Salesforce
    #             if st.button("💾 Opslaan in Salesforce"):
    #                 if selected_account and transcribed_text:
    #                     klantnummer = selected_account.split(" - ")[0]
    #                     save_to_salesforce(sf, klantnummer, transcribed_text)
    #                 else:
    #                     st.warning("⚠️ Selecteer een account en spreek iets in!")

    def genereer_prompt(bedrijfsnaam, vestigingsplaats):
        prompt = (
            f"Neem diep adem, en geef mij een zeer gedetailleerd overzicht van alle beschikbare zakelijke informatie over "
            f"{bedrijfsnaam} te {vestigingsplaats}. Belangrijk vind ik: omzet en andere financiële gegevens, "
            f"producten/diensten, markten waarin ze actief zijn, strategische doelstellingen, aantal werknemers/fte (probeer hier met name via linkedin of kvk te zoeken)."
            f"Check ook op recente overnames of investeringen. "
            f"Focus op informatie die nuttig is voor een verkoopbezoek of voor opname in een CRM-systeem. Benoem de CEO/algemeen directeur/eigenaar en zijn telefoonnummer en e-mailadres. "
            f"Graag samengevat in duidelijke bullets per categorie."
            )
        
        return prompt
    
    def verkrijg_perplexity_response(prompt: str) -> str:
        try:
            headers = {
                "Authorization": f"Bearer {pplx_api_key}",
                "Content-Type": "application/json"
            }
    
            payload = {
                "model": "sonar-pro",
                "messages": [
                    {
                        "role": "system",
                        "content": "Je bent een behulpzame zakelijke assistent. Gebruik bronvermeldingen op een overzichtelijke en professionele manier."
                    },
                    {
                        "role": "user",
                        "content": prompt + (
                            " Gebruik verwijzingen zoals [1], [2], [3] in de hoofdtekst waar nodig. "
                            "Voeg aan het einde van het antwoord een lijst toe met de bijbehorende bronnen in het volgende format:\n"
                            "[1] https://voorbeeld1.nl\n"
                            "[2] https://voorbeeld2.nl\n"
                            "[3] https://voorbeeld3.nl\n"
                            "Zorg ervoor dat de nummers overeenkomen met de verwijzingen in de tekst."
                            "Noem het stuk: 'Scoutverslag van {bedrijfsnaam}'"
                        )
                    }
                ]
            }
                
            response = requests.post("https://api.perplexity.ai/chat/completions", headers=headers, json=payload)
    
            if response.status_code != 200:
                return f"Er is een fout opgetreden: {response.status_code} - {response.text}"
    
            result = response.json()
            antwoord = result["choices"][0]["message"]["content"]
            return antwoord.strip()
    
        except Exception as e:
            return f"Er is een fout opgetreden: {str(e)}"
            
    # Functie om een GPT-resultaat als PDF te genereren
    def generate_pdf_from_text(content: str, bedrijfsnaam: str, vestigingsplaats: str):
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
        from io import BytesIO
    
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=18)
        elements = []
    
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Justify', alignment=TA_LEFT, leading=14))
    
        # Titel
        title = f"Informatie over {bedrijfsnaam} ({vestigingsplaats})"
        elements.append(Paragraph(title, styles['Title']))
        elements.append(Spacer(1, 12))
    
        # GPT-inhoud opsplitsen per regel
        for line in content.split('\n'):
            if line.strip() != "":
                elements.append(Paragraph(line.strip(), styles['Justify']))
                elements.append(Spacer(1, 6))
    
        doc.build(elements)
        pdf = buffer.getvalue()
        buffer.close()
        return pdf
    
    # Streamlit-app lay-out
    st.header("Scout 🕵️‍♂️")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        bedrijfsnaam = st.text_input("Naam van het bedrijf:")
        vestigingsplaats = st.text_input("Vestigingsplaats van het bedrijf:")
    
    if st.button("Stuur scout op pad!"):
        if bedrijfsnaam and vestigingsplaats:
            prompt = genereer_prompt(bedrijfsnaam, vestigingsplaats)
            with st.spinner(f"Aan het scouten..."):
                response = verkrijg_perplexity_response(prompt)
    
            st.markdown("### Resultaten:")
            st.write(response)
    
            pdf_bytes = generate_pdf_from_text(response, bedrijfsnaam, vestigingsplaats)
            st.download_button(
                label="📄 Download resultaat als PDF",
                data=pdf_bytes,
                file_name=f"{bedrijfsnaam.replace(' ', '_')}_informatie.pdf",
                mime="application/pdf"
            )
        else:
            st.warning("Vul zowel de bedrijfsnaam als de vestigingsplaats in.")

    

    # # Ophalen van gegevens
    # if st.button("Haal gegevens op"):
    #     response = session.get(list_items_url, headers=headers)
    #     if response.status_code == 200:
    #         st.success("✅ Gegevens succesvol opgehaald!")
    #         data = response.json()
    #         # Toon de gegevens in een tabel
    #         if "d" in data and "results" in data["d"]:
    #             items = data["d"]["results"]
    #             for item in items:
    #                 st.write(f"ID: {item['Id']}, Synoniem: {item.get('Synoniem', 'N/A')}")
    #         else:
    #             st.warning("⚠️ Geen gegevens gevonden.")
    #     else:
    #         st.error(f"❌ Fout bij ophalen van gegevens: {response.status_code}, {response.text}")
    
# with tab5:
#     st.subheader("💬 Glasadvies Chatbot")
#     st.info("Stel je vraag over glas en krijg advies van AI op basis van beschikbare bronnen.")

    # # Functie om website content en subpagina's op te halen
    # def fetch_website_and_subpages(base_url, max_depth=0):
    #     visited_urls = set()
    #     content_list = []

    #     def crawl(url, depth):
    #         if url in visited_urls or depth > max_depth:
    #             return
    #         try:
    #             visited_urls.add(url)
    #             response = requests.get(url)
    #             if response.status_code == 200:
    #                 soup = BeautifulSoup(response.text, "html.parser")
    #                 content_list.append(soup.get_text())  # Voeg de tekst van de pagina toe
                    
    #                 # Zoek naar onderliggende links
    #                 for link in soup.find_all("a", href=True):
    #                     next_url = link["href"]
    #                     if next_url.startswith("/") or next_url.startswith(base_url):
    #                         full_url = next_url if next_url.startswith("http") else f"{base_url.rstrip('/')}/{next_url.lstrip('/')}"
    #                         crawl(full_url, depth + 1)
    #         except Exception as e:
    #             st.error(f"Fout bij ophalen van {url}: {e}")
    
    #     crawl(base_url, 0)
    #     return "\n".join(content_list)

    # # Functie om PDF-inhoud op te halen
    # def fetch_pdf_content(url):
    #     try:
    #         response = requests.get(url)
    #         pdf_file = io.BytesIO(response.content)
    #         pdf_reader = PdfReader(pdf_file)
    #         text = ""
    #         for page in pdf_reader.pages:
    #             text += page.extract_text()
    #         return text
    #     except Exception as e:
    #         st.error(f"Kon de PDF {url} niet verwerken: {e}")
    #         return ""
    
    # # Bronnen ophalen (websites + PDF)
    # sources = [
    #     fetch_website_and_subpages("https://www.onderhoudnl.nl/glasvraagbaak", max_depth=0),
    #     fetch_website_and_subpages("https://www.glasdiscount.nl/kennisbank/begrippen", max_depth=0),
    #     fetch_pdf_content("https://www.kenniscentrumglas.nl/wp-content/uploads/Infosheet-NEN-2608-1.pdf"),
    #     fetch_pdf_content("https://www.kenniscentrumglas.nl/wp-content/uploads/KCG-infosheet-Letselveiligheid-glas-NEN-3569-1.pdf"),
    # ]
    # combined_source_text = "\n".join(sources)
    
    # Initialiseer chatgeschiedenis in sessiestatus
    # if "chat_history" not in st.session_state:
    #     st.session_state["chat_history"] = [{"role": "assistant", "content": "Hoe kan ik je helpen met glasadvies?"}]
    
    # st.title("💬 Glasadvies Chatbot")
    
    # # Toon chatgeschiedenis
    # for msg in st.session_state["chat_history"]:
    #     st.chat_message(msg["role"]).write(msg["content"])
    
    # # Inputveld voor gebruikersvraag
    # user_query = st.chat_input("Stel je vraag hier:")
    
    # if user_query:
    #     st.chat_message("user").write(user_query)  # Toon de gebruikersvraag
    #     st.session_state["chat_history"].append({"role": "user", "content": user_query})
    
    #     try:
    #         # # Verstuur de vraag naar OpenAI met de opgehaalde documentatie
    #         # response = openai.chat.completions.create(
    #         #     model="gpt-4",
    #         #     messages=[
    #         #         {"role": "system", "content": "Je bent een glasadvies assistent die technisch advies geeft op basis van de gegeven documentatie. Geef kort en helder advies."},
    #         #         {"role": "user", "content": f"Documentatie:\n{combined_source_text}\n\nVraag: {user_query}"}
    #         #     ],
    #         #     max_tokens=300,
    #         #     temperature=0.7
    #         # )

    #         # # Toon het antwoord van OpenAI
    #         # ai_response = response.choices[0].message.content
    #         # st.chat_message("assistant").write(ai_response)
    #         # st.session_state["chat_history"].append({"role": "assistant", "content": ai_response})
    #         pass  # Deze logica wordt niet uitgevoerd
    #     except Exception as e:
    #         st.error(f"Er is een fout opgetreden bij het raadplegen van OpenAI: {e}")

            
